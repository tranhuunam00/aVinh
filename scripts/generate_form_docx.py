import os
import sys
import docx
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
import pymupdf
import win32com.client
import zipfile

sys.stdout.reconfigure(encoding='utf-8')

base_dir = r"d:\DAOG\aVinh"
temp_docx_path = os.path.join(base_dir, "worddata", "temp_build.docx")
final_docx_path = os.path.join(base_dir, "worddata", "Phieu_Dat_Hang_Nhiem_Vu_KHCN_Mau_I.01-DHNV.docx")
pdf_path = os.path.join(base_dir, "worddata", "Phieu_Dat_Hang_Nhiem_Vu_KHCN_Mau_I.01-DHNV.pdf")

doc = docx.Document()

# 1. Page Setup - A4 (210 x 297 mm)
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin = Mm(30)
section.right_margin = Mm(18)
section.header_distance = Mm(12)
section.footer_distance = Mm(12)

# Normal Style
normal_style = doc.styles['Normal']
normal_style.font.name = 'Times New Roman'
normal_style.font.size = Pt(14)
normal_style.font.color.rgb = RGBColor(0, 0, 0)
normal_style.paragraph_format.space_before = Pt(0)
normal_style.paragraph_format.space_after = Pt(4)
normal_style.paragraph_format.line_spacing = 1.2

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_p(doc, text="", bold=False, italic=False, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, line_spacing=1.2, keep_with_next=False, has_tab_dots=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.keep_with_next = keep_with_next
    if has_tab_dots:
        p.paragraph_format.tab_stops.add_tab_stop(Mm(162), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    if text:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p

# 1. Header: PHỤ LỤC II. BIỂU MẪU
add_p(doc, "PHỤ LỤC II. BIỂU MẪU", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2, line_spacing=1.15)

# 2. Header: Mẫu I.01-ĐHNV
add_p(doc, "Mẫu I.01-ĐHNV", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=4, line_spacing=1.15)

# 3. Top Organ and Motto Table (2 columns, borderless)
top_table = doc.add_table(rows=1, cols=2)
top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
top_table.autofit = False

cell_left = top_table.cell(0, 0)
cell_right = top_table.cell(0, 1)

cell_left.width = Mm(52)
cell_right.width = Mm(110)

# Left cell: Organ
set_cell_margins(cell_left, top=0, bottom=0, left=0, right=0)
p_org1 = cell_left.paragraphs[0]
p_org1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_org1.paragraph_format.space_before = Pt(0)
p_org1.paragraph_format.space_after = Pt(1)
p_org1.paragraph_format.line_spacing = 1.15
r = p_org1.add_run("TÊN CƠ QUAN CHỦ QUẢN")
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
r.font.bold = False

p_org2 = cell_left.add_paragraph()
p_org2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_org2.paragraph_format.space_before = Pt(0)
p_org2.paragraph_format.space_after = Pt(0)
p_org2.paragraph_format.line_spacing = 1.15
r = p_org2.add_run("TÊN TỔ CHỨC")
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
r.font.bold = True

p_org3 = cell_left.add_paragraph()
p_org3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_org3.paragraph_format.space_before = Pt(0)
p_org3.paragraph_format.space_after = Pt(0)
p_org3.paragraph_format.line_spacing = 1.0
r = p_org3.add_run("______")
r.font.name = 'Times New Roman'
r.font.size = Pt(11)
r.font.bold = True

# Right cell: Motto & Date
set_cell_margins(cell_right, top=0, bottom=0, left=0, right=0)
p_motto1 = cell_right.paragraphs[0]
p_motto1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_motto1.paragraph_format.space_before = Pt(0)
p_motto1.paragraph_format.space_after = Pt(1)
p_motto1.paragraph_format.line_spacing = 1.15
r = p_motto1.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
r.font.bold = True

p_motto2 = cell_right.add_paragraph()
p_motto2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_motto2.paragraph_format.space_before = Pt(0)
p_motto2.paragraph_format.space_after = Pt(0)
p_motto2.paragraph_format.line_spacing = 1.15
r = p_motto2.add_run("Độc lập - Tự do - Hạnh phúc")
r.font.name = 'Times New Roman'
r.font.size = Pt(13)
r.font.bold = True

p_motto3 = cell_right.add_paragraph()
p_motto3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_motto3.paragraph_format.space_before = Pt(0)
p_motto3.paragraph_format.space_after = Pt(3)
p_motto3.paragraph_format.line_spacing = 1.0
r = p_motto3.add_run("_________________")
r.font.name = 'Times New Roman'
r.font.size = Pt(11)
r.font.bold = True

p_date = cell_right.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date.paragraph_format.space_before = Pt(0)
p_date.paragraph_format.space_after = Pt(0)
p_date.paragraph_format.line_spacing = 1.15
r_dots = p_date.add_run("........")
r_dots.font.name = 'Times New Roman'
r_dots.font.size = Pt(13)
r_dots.font.italic = False

r_dtext = p_date.add_run(", ngày… tháng…  năm 20....")
r_dtext.font.name = 'Times New Roman'
r_dtext.font.size = Pt(13)
r_dtext.font.italic = True

# 4. Title: PHIẾU ĐẶT HÀNG NHIỆM VỤ KHOA HỌC VÀ CÔNG NGHỆ
p_title = add_p(doc, "", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=14, space_after=14, line_spacing=1.15)
r_title = p_title.add_run("PHIẾU ĐẶT HÀNG NHIỆM VỤ KHOA HỌC VÀ CÔNG NGHỆ")
r_title.font.name = 'Times New Roman'
r_title.font.size = Pt(14)
r_title.font.bold = True

# Add Footnote Reference XML with SUPERSCRIPT vertical alignment and larger font size (w:sz=24 -> 12pt)
fn_ref = parse_xml(
    r'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    r'<w:rPr>'
    r'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
    r'<w:rStyle w:val="FootnoteReference"/>'
    r'<w:vertAlign w:val="superscript"/>'
    r'<w:b/>'
    r'<w:sz w:val="24"/>'
    r'</w:rPr>'
    r'<w:footnoteReference w:id="1"/>'
    r'</w:r>'
)
p_title._p.append(fn_ref)

# 5. Body Items (Page 1)
add_p(doc, "1. Tên nhiệm vụ:\t", space_after=2, has_tab_dots=True)
add_p(doc, "\t", space_after=4, has_tab_dots=True)

add_p(doc, "2. Thuộc lĩnh vực:", space_after=2)

# Checkbox Table for Fields (4 rows x 2 cols)
tbl_field = doc.add_table(rows=4, cols=2)
tbl_field.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_field.autofit = False

fields_data = [
    ("- Khoa học tự nhiên", "- Khoa học kỹ thuật và công nghệ"),
    ("- Khoa học y, dược", "- Khoa học nông nghiệp"),
    ("- Khoa học xã hội", "- Khoa học nhân văn"),
    ("- Công nghệ chiến lược", "")
]

for row_idx, (col1_txt, col2_txt) in enumerate(fields_data):
    cell1 = tbl_field.cell(row_idx, 0)
    cell2 = tbl_field.cell(row_idx, 1)
    cell1.width = Mm(81)
    cell2.width = Mm(81)
    for col_idx, (cell, txt) in enumerate([(cell1, col1_txt), (cell2, col2_txt)]):
        set_cell_margins(cell, top=0, bottom=0, left=50, right=50)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.tab_stops.add_tab_stop(Mm(76), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        if txt:
            r1 = p.add_run(txt + "\t")
            r1.font.name = 'Times New Roman'
            r1.font.size = Pt(14)
            r2 = p.add_run("☐")
            r2.font.name = 'Segoe UI Symbol'
            r2.font.size = Pt(14)

add_p(doc, "3. Thuộc loại hình nhiệm vụ", space_before=4, space_after=2)

# Checkbox Table for Types (2 rows x 2 cols)
tbl_type = doc.add_table(rows=2, cols=2)
tbl_type.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_type.autofit = False

types_data = [
    ("- Nghiên cứu cơ bản", "- Nghiên cứu ứng dụng"),
    ("- Phát triển công nghệ", "- Phát triển giải pháp xã hội")
]

for row_idx, (col1_txt, col2_txt) in enumerate(types_data):
    cell1 = tbl_type.cell(row_idx, 0)
    cell2 = tbl_type.cell(row_idx, 1)
    cell1.width = Mm(81)
    cell2.width = Mm(81)
    for col_idx, (cell, txt) in enumerate([(cell1, col1_txt), (cell2, col2_txt)]):
        set_cell_margins(cell, top=0, bottom=0, left=50, right=50)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.tab_stops.add_tab_stop(Mm(76), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        r1 = p.add_run(txt + "\t")
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(14)
        r2 = p.add_run("☐")
        r2.font.name = 'Segoe UI Symbol'
        r2.font.size = Pt(14)

add_p(doc, "4. Căn cứ đặt hàng nhiệm vụ:\t", space_before=4, space_after=2, has_tab_dots=True)
add_p(doc, "\t", space_after=4, has_tab_dots=True)

add_p(doc, "5. Tính cấp thiết của nhiệm vụ:\t", space_after=2, has_tab_dots=True)
add_p(doc, "\t", space_after=4, has_tab_dots=True)

add_p(doc, "6. Mục tiêu:\t", space_after=2, has_tab_dots=True)
add_p(doc, "\t", space_after=4, has_tab_dots=True)

add_p(doc, "7. Dự kiến các nội dung chính cần thực hiện:\t", space_after=2, has_tab_dots=True)
add_p(doc, "\t", space_after=4, has_tab_dots=True)

add_p(doc, "8. Dự kiến các kết quả thực hiện nhiệm vụ và các chỉ tiêu cần đạt:\t", space_after=2, has_tab_dots=True)
add_p(doc, "\t", space_after=4, has_tab_dots=True)

add_p(doc, "9. Hiệu quả và tác động của kết quả thực hiện nhiệm vụ:\t", space_after=2, has_tab_dots=True)
add_p(doc, "\t", space_after=4, has_tab_dots=True)

# Items for Page 2
add_p(doc, "10. Dự kiến kinh phí thực hiện:\t", space_after=2, keep_with_next=True, has_tab_dots=True)
add_p(doc, "- Kinh phí hỗ trợ từ NSNN:\t", space_after=2, keep_with_next=True, has_tab_dots=True)
add_p(doc, "- Kinh phí từ nguồn khác:\t", space_after=3, has_tab_dots=True)
add_p(doc, "11. Dự kiến thời gian thực hiện:\t", space_after=4, has_tab_dots=True)

add_p(doc, "12. Đề xuất tổ chức được xét giao trực tiếp (nếu có):\t", space_after=2, has_tab_dots=True)
add_p(doc, "\t", space_after=4, has_tab_dots=True)

add_p(doc, "13. Tiếp nhận và phương án tổ chức quản lý, sử dụng kết quả của nhiệm vụ của", space_after=1)
add_p(doc, "cơ quan đặt hàng: ", space_after=2)
add_p(doc, "(Nêu rõ việc Nhà nước sẽ nắm giữ quyền quản lý, sử dụng, quyền sở hữu kết quả trong trường hợp nhà nước có yêu cầu tiếp nhận kết quả)", italic=True, space_after=2)
add_p(doc, "\t", space_after=2, has_tab_dots=True)
add_p(doc, "\t", space_after=4, has_tab_dots=True)

add_p(doc, "14. Danh mục tài liệu tham khảo:", space_after=4)

add_p(doc, "15. Thông tin liên hệ:", space_after=2)
add_p(doc, "Tên tổ chức:\t", space_after=2, has_tab_dots=True)
add_p(doc, "Đại diện tổ chức:\t", space_after=2, has_tab_dots=True)
add_p(doc, "Mã định danh điện tử của tổ chức:\t", space_after=2, has_tab_dots=True)
add_p(doc, "Điện thoại:\t", space_after=2, has_tab_dots=True)
add_p(doc, "Email:\t", space_after=2, has_tab_dots=True)
add_p(doc, "Địa chỉ liên hệ:\t", space_after=14, has_tab_dots=True)

# Signature block (Table 1 row x 2 cols)
tbl_sig = doc.add_table(rows=1, cols=2)
tbl_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_sig.autofit = False

cell_sig_l = tbl_sig.cell(0, 0)
cell_sig_r = tbl_sig.cell(0, 1)
cell_sig_l.width = Mm(70)
cell_sig_r.width = Mm(92)

set_cell_margins(cell_sig_r, top=0, bottom=0, left=0, right=0)
p_s1 = cell_sig_r.paragraphs[0]
p_s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_s1.paragraph_format.space_before = Pt(8)
p_s1.paragraph_format.space_after = Pt(2)
p_s1.paragraph_format.line_spacing = 1.15
r = p_s1.add_run("ĐẠI DIỆN TỔ CHỨC ĐẶT HÀNG")
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.font.bold = True

p_s2 = cell_sig_r.add_paragraph()
p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_s2.paragraph_format.space_before = Pt(0)
p_s2.paragraph_format.space_after = Pt(0)
p_s2.paragraph_format.line_spacing = 1.15
r = p_s2.add_run("(Họ, tên và chữ ký - đóng dấu)")
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.font.italic = True

# Save temporary docx
doc.save(temp_docx_path)

# Inject footnote XML
def inject_footnote_into_docx(in_docx_path, out_docx_path, footnote_text):
    footnotes_xml_content = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:footnote w:type="separator" w:id="-1">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:id="1">
    <w:p>
      <w:pPr>
        <w:pStyle w:val="FootnoteText"/>
        <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
          <w:rStyle w:val="FootnoteReference"/>
          <w:vertAlign w:val="superscript"/>
          <w:sz w:val="22"/>
        </w:rPr>
        <w:footnoteRef/>
      </w:r>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
          <w:i/>
          <w:sz w:val="24"/>
        </w:rPr>
        <w:t xml:space="preserve"> {footnote_text}</w:t>
      </w:r>
    </w:p>
  </w:footnote>
</w:footnotes>""".encode('utf-8')

    with zipfile.ZipFile(in_docx_path, 'r') as zin:
        with zipfile.ZipFile(out_docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    ct_str = data.decode('utf-8')
                    if 'footnotes.xml' not in ct_str:
                        fn_override = '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                        ct_str = ct_str.replace('</Types>', f'{fn_override}</Types>')
                        data = ct_str.encode('utf-8')
                elif item.filename == 'word/_rels/document.xml.rels':
                    rels_str = data.decode('utf-8')
                    if 'footnotes.xml' not in rels_str:
                        fn_rel = '<Relationship Id="rIdFootnotes" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
                        rels_str = rels_str.replace('</Relationships>', f'{fn_rel}</Relationships>')
                        data = rels_str.encode('utf-8')
                zout.writestr(item, data)
            zout.writestr('word/footnotes.xml', footnotes_xml_content)

inject_footnote_into_docx(temp_docx_path, final_docx_path, "Phiếu đặt hàng được trình bày không quá 10 trang giấy khổ A4.")
print("Saved final DOCX:", final_docx_path)

# Convert to PDF via Word COM
word = win32com.client.Dispatch("Word.Application")
word.Visible = False
word.DisplayAlerts = 0
try:
    doc_word = word.Documents.Open(os.path.abspath(final_docx_path))
    doc_word.ExportAsFixedFormat(os.path.abspath(pdf_path), 17)
    print("Exported PDF:", pdf_path)
    doc_word.Close(False)
finally:
    word.Quit()

# Render PDF to PNGs to inspect
pdf_doc = pymupdf.open(pdf_path)
print(f"Generated PDF has {len(pdf_doc)} pages")
for p_idx, page in enumerate(pdf_doc):
    pix = page.get_pixmap(dpi=200)
    out_img = os.path.join(base_dir, "worddata", f"gen_page_{p_idx+1}.png")
    pix.save(out_img)
    print(f"Rendered: {out_img}")
