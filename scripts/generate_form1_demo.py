import os
import sys
import zipfile
import docx
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import win32com.client
import pymupdf
import time
import gc

sys.stdout.reconfigure(encoding='utf-8')

BASE_DIR = r"d:\DAOGROUP_WORKSPACE\aVinh"
OUT_DIR = os.path.join(BASE_DIR, "worddata", "chinhthuc")
os.makedirs(OUT_DIR, exist_ok=True)

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_p(doc, text="", bold=False, italic=False, font_size=12.5, align=WD_ALIGN_PARAGRAPH.LEFT, 
          space_before=0, space_after=1.5, line_spacing=1.12, first_line_indent_mm=0, left_indent_mm=0, keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.keep_with_next = keep_with_next
    if first_line_indent_mm > 0:
        p.paragraph_format.first_line_indent = Mm(first_line_indent_mm)
    if left_indent_mm > 0:
        p.paragraph_format.left_indent = Mm(left_indent_mm)
    if text:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p

def create_base_document(left_mm=28, right_mm=18, top_mm=18, bottom_mm=18):
    doc = docx.Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Mm(top_mm)
    section.bottom_margin = Mm(bottom_mm)
    section.left_margin = Mm(left_mm)
    section.right_margin = Mm(right_mm)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12.5)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(1.5)
    normal_style.paragraph_format.line_spacing = 1.12
    return doc

def inject_footnotes_into_docx(in_docx_path, out_docx_path, footnotes_dict):
    fn_items = ""
    for fn_id, fn_text in footnotes_dict.items():
        fn_items += f"""
  <w:footnote w:id="{fn_id}">
    <w:p>
      <w:pPr>
        <w:pStyle w:val="FootnoteText"/>
        <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
          <w:rStyle w:val="FootnoteReference"/>
          <w:vertAlign w:val="superscript"/>
          <w:sz w:val="20"/>
        </w:rPr>
        <w:footnoteRef/>
      </w:r>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
          <w:i/>
          <w:sz w:val="22"/>
        </w:rPr>
        <w:t xml:space="preserve"> {fn_text}</w:t>
      </w:r>
    </w:p>
  </w:footnote>"""

    footnotes_xml_content = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:footnote w:type="separator" w:id="-1">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>{fn_items}
</w:footnotes>""".encode('utf-8')

    staging_path = out_docx_path + ".tmp"
    with zipfile.ZipFile(in_docx_path, 'r') as zin:
        with zipfile.ZipFile(staging_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    ct_str = data.decode('utf-8')
                    if 'footnotes.xml' not in ct_str:
                        fn_override = '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                        ct_str = ct_str.replace('</Types>', f'{fn_override}</Types>')
                        data = ct_str.encode('utf-8')
                elif item.filename == 'word/_rels/document.xml.rels':
                    rels_str = data.decode('utf-8')
                    if 'footnotes.xml' not in rels_str:
                        fn_rel = '<Relationship Id="rIdFootnotes" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
                        rels_str = rels_str.replace('</Relationships>', f'{fn_rel}</Relationships>')
                        data = rels_str.encode('utf-8')
                zout.writestr(item, data)
            zout.writestr('word/footnotes.xml', footnotes_xml_content)

    for _ in range(5):
        try:
            if os.path.exists(out_docx_path):
                try:
                    os.remove(out_docx_path)
                except Exception:
                    pass
            os.replace(staging_path, out_docx_path)
            break
        except Exception:
            time.sleep(0.5)

def export_docx_to_pdf(docx_path, pdf_path):
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc_word = None
    try:
        doc_word = word.Documents.Open(os.path.abspath(docx_path))
        doc_word.ExportAsFixedFormat(os.path.abspath(pdf_path), 17)
        doc_word.Close(False)
        doc_word = None
        print(f"  [PDF Exported] {pdf_path}")
    finally:
        word.Quit()
        if doc_word is not None:
            try:
                doc_word.Close(False)
            except Exception:
                pass
            del doc_word
        del word
        gc.collect()
        time.sleep(0.3)

print("================================================================================")
print("BẮT ĐẦU TẠO DỮ LIỆU DEMO AN TOÀN VÀ ĐIỀN VÀO FORM 1 (PHIẾU ĐẶT HÀNG NHIỆM VỤ KH&CN)")
print("================================================================================")

doc = create_base_document(left_mm=28, right_mm=18, top_mm=18, bottom_mm=18)

# 1. Header Tiêu ngữ
add_p(doc, "PHỤ LỤC II. BIỂU MẪU", bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1, line_spacing=1.1)
add_p(doc, "Mẫu I.01-ĐHNV", bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=3, line_spacing=1.1)

top_table = doc.add_table(rows=1, cols=2)
top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
top_table.autofit = False

c_left = top_table.cell(0, 0)
c_right = top_table.cell(0, 1)
c_left.width = Mm(72)
c_right.width = Mm(92)

set_cell_margins(c_left, top=0, bottom=0, left=0, right=0)
p_org1 = c_left.paragraphs[0]
p_org1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_org1.paragraph_format.space_before, p_org1.paragraph_format.space_after, p_org1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.1
r = p_org1.add_run("ỦY BAN NHÂN DÂN TỈNH ĐỒNG NAI")
r.font.size = Pt(11)

p_org2 = c_left.add_paragraph()
p_org2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_org2.paragraph_format.space_before, p_org2.paragraph_format.space_after, p_org2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.1
r = p_org2.add_run("SỞ Y TẾ TỈNH ĐỒNG NAI")
r.font.size, r.font.bold = Pt(11.5), True

p_org3 = c_left.add_paragraph()
p_org3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_org3.paragraph_format.space_before, p_org3.paragraph_format.space_after, p_org3.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
r = p_org3.add_run("______")
r.font.size, r.font.bold = Pt(10), True

set_cell_margins(c_right, top=0, bottom=0, left=0, right=0)
p_m1 = c_right.paragraphs[0]
p_m1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_m1.paragraph_format.space_before, p_m1.paragraph_format.space_after, p_m1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.1
r = p_m1.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
r.font.size, r.font.bold = Pt(11), True

p_m2 = c_right.add_paragraph()
p_m2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_m2.paragraph_format.space_before, p_m2.paragraph_format.space_after, p_m2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.1
r = p_m2.add_run("Độc lập - Tự do - Hạnh phúc")
r.font.size, r.font.bold = Pt(12), True

p_m3 = c_right.add_paragraph()
p_m3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_m3.paragraph_format.space_before, p_m3.paragraph_format.space_after, p_m3.paragraph_format.line_spacing = Pt(0), Pt(2), 1.0
r = p_m3.add_run("_________________")
r.font.size, r.font.bold = Pt(10), True

p_d = c_right.add_paragraph()
p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_d.paragraph_format.space_before, p_d.paragraph_format.space_after, p_d.paragraph_format.line_spacing = Pt(0), Pt(0), 1.1
r_d2 = p_d.add_run("Đồng Nai, ngày 25 tháng 08 năm 2026")
r_d2.font.size, r_d2.font.italic = Pt(12), True

# Title: PHIẾU ĐẶT HÀNG NHIỆM VỤ KHOA HỌC VÀ CÔNG NGHỆ
p_title = add_p(doc, "", bold=True, font_size=13.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10, line_spacing=1.15)
r_t = p_title.add_run("PHIẾU ĐẶT HÀNG NHIỆM VỤ KHOA HỌC VÀ CÔNG NGHỆ")
r_t.font.size, r_t.font.bold = Pt(13.5), True

fn_ref = parse_xml(
    f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
    f'<w:rStyle w:val="FootnoteReference"/><w:vertAlign w:val="superscript"/><w:b/><w:sz w:val="22"/></w:rPr>'
    f'<w:footnoteReference w:id="1"/></w:r>'
)
p_title._p.append(fn_ref)

# --- DỮ LIỆU ĐIỀN FORM (AN TOÀN / DEMO TRÌNH CHIẾU) ---

# 1. Tên nhiệm vụ
p1 = add_p(doc, "1. Tên nhiệm vụ: ", bold=True, font_size=12.5, space_after=2)
r = p1.add_run("Nghiên cứu ứng dụng trí tuệ nhân tạo (AI) và công nghệ sinh học tiên tiến (huyết tương giàu tiểu cầu PRP tự thân, phẫu thuật ít xâm lấn MIS) trong sàng lọc sớm, chẩn đoán và điều trị phục hồi thoái hóa khớp gối tại cộng đồng tỉnh Đồng Nai.")
r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

# 2. Thuộc lĩnh vực
add_p(doc, "2. Thuộc lĩnh vực:", bold=True, font_size=12.5, space_after=1.5)
tbl_field = doc.add_table(rows=4, cols=2)
tbl_field.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_field.autofit = False

fields_data = [
    ("- Khoa học tự nhiên", False, "- Khoa học kỹ thuật và công nghệ", False),
    ("- Khoa học y, dược", True, "- Khoa học nông nghiệp", False),
    ("- Khoa học xã hội", False, "- Khoa học nhân văn", False),
    ("- Công nghệ chiến lược", False, "", False)
]
for row_idx, (t1, c1, t2, c2) in enumerate(fields_data):
    cell1, cell2 = tbl_field.cell(row_idx, 0), tbl_field.cell(row_idx, 1)
    cell1.width, cell2.width = Mm(82), Mm(82)
    for cell, txt, is_checked in [(cell1, t1, c1), (cell2, t2, c2)]:
        set_cell_margins(cell, top=0, bottom=0, left=40, right=40)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(1), 1.12
        p.paragraph_format.tab_stops.add_tab_stop(Mm(76), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        if txt:
            r1 = p.add_run(txt + "\t")
            r1.font.size = Pt(12)
            r2 = p.add_run("☒" if is_checked else "☐")
            r2.font.name, r2.font.size = 'Segoe UI Symbol', Pt(12)
            if is_checked:
                r2.font.bold = True

# 3. Thuộc loại hình nhiệm vụ
add_p(doc, "3. Thuộc loại hình nhiệm vụ", bold=True, font_size=12.5, space_before=2, space_after=1.5)
tbl_type = doc.add_table(rows=2, cols=2)
tbl_type.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_type.autofit = False

types_data = [
    ("- Nghiên cứu cơ bản", False, "- Nghiên cứu ứng dụng", True),
    ("- Phát triển công nghệ", False, "- Phát triển giải pháp xã hội", False)
]
for row_idx, (t1, c1, t2, c2) in enumerate(types_data):
    cell1, cell2 = tbl_type.cell(row_idx, 0), tbl_type.cell(row_idx, 1)
    cell1.width, cell2.width = Mm(82), Mm(82)
    for cell, txt, is_checked in [(cell1, t1, c1), (cell2, t2, c2)]:
        set_cell_margins(cell, top=0, bottom=0, left=40, right=40)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(1), 1.12
        p.paragraph_format.tab_stops.add_tab_stop(Mm(76), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        r1 = p.add_run(txt + "\t")
        r1.font.size = Pt(12)
        r2 = p.add_run("☒" if is_checked else "☐")
        r2.font.name, r2.font.size = 'Segoe UI Symbol', Pt(12)
        if is_checked:
            r2.font.bold = True

# 4. Căn cứ đặt hàng nhiệm vụ
add_p(doc, "4. Căn cứ đặt hàng nhiệm vụ:", bold=True, font_size=12.5, space_before=2, space_after=1.5)
add_p(doc, "- Nghị quyết số 57-NQ/TW của Bộ Chính trị về phát triển và ứng dụng công nghệ sinh học và kỹ thuật y học cao phục vụ chăm sóc sức khỏe nhân dân;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "- Quyết định số 18/2026/QĐ-UBND của UBND tỉnh Đồng Nai quy định quản lý nhiệm vụ khoa học công nghệ và đổi mới sáng tạo cấp tỉnh;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "- Kế hoạch phát triển công nghệ cao và y tế thông minh tỉnh Đồng Nai giai đoạn 2026 - 2030 (Kế hoạch số 115/KH-UBND của UBND tỉnh).", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=10)

# 5. Tính cấp thiết
add_p(doc, "5. Tính cấp thiết của nhiệm vụ:", bold=True, font_size=12.5, space_after=1.5)
add_p(doc, "Thoái hóa khớp gối là nguyên nhân hàng đầu gây tàn phế vận động ở người cao tuổi và lao động trung niên. Tại địa bàn tỉnh, tỷ lệ mắc bệnh có xu hướng trẻ hóa và gia tăng nhanh nhưng công tác phát hiện sớm còn hạn chế. Đa số bệnh nhân chỉ đến khám khi sụn khớp đã tổn thương nặng (giai đoạn III-IV), dẫn đến chi phí điều trị lớn và hiệu quả phục hồi kém. Việc triển khai nhiệm vụ khoa học này sẽ giúp giải quyết đồng bộ từ khâu sàng lọc dịch tễ bằng công cụ AI hỗ trợ đến phác đồ bảo tồn sinh học tiên tiến (PRP) và phẫu thuật ít xâm lấn MIS, nâng cao vượt bậc chất lượng sống cho người dân ngay tại y tế địa phương.", 
      font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=12.7)

# 6. Mục tiêu
add_p(doc, "6. Mục tiêu:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
add_p(doc, "- Mục tiêu chung: Xây dựng cơ sở dữ liệu dịch tễ học và hoàn thiện mô hình ứng dụng công nghệ y học cao (PRP, MIS, Robot/AI phục hồi chức năng) trong quản lý, chẩn đoán và điều trị thoái hóa khớp gối trên địa bàn tỉnh Đồng Nai.", 
      font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "- Mục tiêu cụ thể: (1) Khám sàng lọc dịch tễ cho 2.000 người dân từ 40 tuổi trở lên; (2) Chuẩn hóa quy trình tiêm PRP tự thân và phẫu thuật nội soi ít xâm lấn; (3) Hoàn thiện 01 phần mềm AI hỗ trợ phân tích hình ảnh X-quang khớp gối; (4) Chuyển giao gói kỹ thuật điều trị bảo tồn cho 3-5 bệnh viện tuyến huyện.", 
      font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=10)

# 7. Dự kiến các nội dung chính cần thực hiện
add_p(doc, "7. Dự kiến các nội dung chính cần thực hiện:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
add_p(doc, "• Nội dung 1: Xây dựng bộ công cụ dịch tễ, quy trình sàng lọc và giao thức nghiên cứu lâm sàng;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "• Nội dung 2: Khám sàng lọc cộng đồng (n = 2.000), chẩn đoán cận lâm sàng chuyên sâu (X-quang, MRI) và xây dựng cơ sở dữ liệu dịch tễ học;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "• Nội dung 3: Phát triển mô hình AI hỗ trợ phân loại mức độ thoái hóa khớp theo Kellgren-Lawrence;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "• Nội dung 4: Thử nghiệm lâm sàng chuỗi can thiệp công nghệ cao (Liệu pháp PRP tự thân kết hợp MIS và bài tập robot thông minh);", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "• Nội dung 5: Đánh giá theo dõi dọc hiệu quả sau can thiệp (1, 3, 6, 12 tháng) qua các chỉ số VAS, WOMAC;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "• Nội dung 6: Xây dựng bộ tài liệu hướng dẫn chuyên môn và tập huấn chuyển giao cho các cơ sở y tế tuyến dưới.", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=10)

# 8. Dự kiến các kết quả thực hiện nhiệm vụ và các chỉ tiêu cần đạt
add_p(doc, "8. Dự kiến các kết quả thực hiện nhiệm vụ và các chỉ tiêu cần đạt:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
add_p(doc, "1. Báo cáo khoa học phân tích dịch tễ học thoái hóa khớp gối tại tỉnh Đồng Nai (n = 2.000 mẫu);", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "2. 03 Quy trình kỹ thuật chuyên môn được phê duyệt (Quy trình tách chiết & tiêm PRP; Quy trình phẫu thuật MIS; Quy trình phục hồi chức năng tích hợp AI);", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "3. 01 Module phần mềm AI phân tích hình ảnh X-quang khớp gối (độ chính xác > 88%);", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "4. 02 Bài báo khoa học đăng trên tạp chí y học chuyên ngành uy tín;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "5. Đào tạo, cấp chứng nhận chuyển giao kỹ thuật cho 50 bác sĩ và kỹ thuật viên tuyến huyện.", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=10)

# 9. Hiệu quả và tác động
add_p(doc, "9. Hiệu quả và tác động của kết quả thực hiện nhiệm vụ:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
add_p(doc, "Kết quả nghiên cứu giúp giảm 25-30% tỷ lệ bệnh nhân phải phẫu thuật thay khớp nhân tạo tốn kém thông qua can thiệp bảo tồn sớm bằng PRP; rút ngắn thời gian phục hồi chức năng vận động từ 6 tháng xuống còn 2-3 tháng nhờ hỗ trợ của AI và Robot; đồng thời nâng cao năng lực chẩn đoán và điều trị của hệ thống y tế công lập tuyến cơ sở, tiết kiệm hàng chục tỷ đồng chi phí điều trị và an sinh xã hội cho người dân địa phương.", 
      font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3, first_line_indent_mm=12.7)

# 10. Dự kiến kinh phí thực hiện
p10 = add_p(doc, "10. Dự kiến kinh phí thực hiện: ", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
r = p10.add_run("16.800.000.000 đồng (Mười sáu tỷ tám trăm triệu đồng chẵn).")
r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

# 11. Dự kiến thời gian thực hiện
p11 = add_p(doc, "11. Dự kiến thời gian thực hiện: ", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
r = p11.add_run("24 tháng (Từ tháng 01 năm 2027 đến tháng 12 năm 2028).")
r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

# Sub-bullets of kinh phí
p_ns = add_p(doc, "- Kinh phí hỗ trợ từ NSNN: ", bold=True, font_size=12.5, space_after=1.5, left_indent_mm=8, keep_with_next=True)
r = p_ns.add_run("4.200.000.000 đồng (Bốn tỷ hai trăm triệu đồng chẵn).")
r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

p_nk = add_p(doc, "- Kinh phí từ nguồn khác: ", bold=True, font_size=12.5, space_after=2.5, left_indent_mm=8, keep_with_next=True)
r = p_nk.add_run("12.600.000.000 đồng (Nguồn vốn đối ứng tự có của đơn vị chủ trì thực hiện).")
r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

# 12. Đề xuất tổ chức được xét giao trực tiếp
p12 = add_p(doc, "12. Đề xuất tổ chức được xét giao trực tiếp (nếu có): ", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
r = p12.add_run("Viện Nghiên cứu và Ứng dụng Y học Kỹ thuật cao Sao Mai (Địa chỉ: Khu Đô thị Công nghệ cao, Tỉnh Đồng Nai).")
r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

# 13. Tiếp nhận và phương án quản lý
add_p(doc, "13. Tiếp nhận và phương án tổ chức quản lý, sử dụng kết quả của nhiệm vụ của cơ quan đặt hàng:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
add_p(doc, "(Nêu rõ việc Nhà nước sẽ nắm giữ quyền quản lý, sử dụng, quyền sở hữu kết quả trong trường hợp nhà nước có yêu cầu tiếp nhận kết quả)", italic=True, font_size=11.5, space_after=1.5, left_indent_mm=5)
add_p(doc, "- Cơ quan đặt hàng (Sở Y tế tỉnh Đồng Nai) là đơn vị đại diện chủ sở hữu, tiếp nhận bàn giao toàn bộ báo cáo dịch tễ học, 03 bộ quy trình kỹ thuật chuẩn và module phần mềm AI sau khi nghiệm thu nhiệm vụ;", 
      font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "- Sở Y tế sẽ ban hành Hướng dẫn chuyên môn áp dụng thống nhất trên toàn tỉnh; giao các bệnh viện đa khoa tuyến huyện tổ chức tiếp nhận chuyển giao quy trình kỹ thuật tiêm PRP và phác đồ phục hồi chức năng thông minh để phục vụ trực tiếp người dân;", 
      font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
add_p(doc, "- Nhà nước nắm giữ toàn quyền quản lý, sử dụng kết quả nghiên cứu và quyền sở hữu đối với các sản phẩm được tài trợ từ nguồn ngân sách KH&CN cấp tỉnh theo đúng quy định hiện hành.", 
      font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=10)

# 14. Danh mục tài liệu tham khảo
add_p(doc, "14. Danh mục tài liệu tham khảo:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
add_p(doc, "[1] Bộ Y tế (2020), Hướng dẫn chẩn đoán và điều trị các bệnh về cơ xương khớp, NXB Y học, Hà Nội.", font_size=11.5, space_after=1, first_line_indent_mm=10)
add_p(doc, "[2] OARSI (2019), Guidelines for the non-surgical management of knee osteoarthritis, Osteoarthritis and Cartilage, 27(11), pp. 1578-1589.", font_size=11.5, space_after=1, first_line_indent_mm=10)
add_p(doc, "[3] American College of Rheumatology (2020), Guideline for the Management of Osteoarthritis of the Hand, Hip, and Knee, Arthritis Care & Research, 72(2), pp. 149-162.", font_size=11.5, space_after=1, first_line_indent_mm=10)
add_p(doc, "[4] Filardo G. et al. (2021), Platelet-rich plasma intra-articular injections for osteoarthritis: A systematic review, Am J Sports Med, 49(2), pp. 529-537.", font_size=11.5, space_after=2.5, first_line_indent_mm=10)

# 15. Thông tin liên hệ
add_p(doc, "15. Thông tin liên hệ:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
add_p(doc, "Tên tổ chức: SỞ Y TẾ TỈNH ĐỒNG NAI", font_size=12.5, space_after=1, left_indent_mm=10, keep_with_next=True)
add_p(doc, "Đại diện tổ chức: Lãnh đạo Sở Y tế tỉnh Đồng Nai", font_size=12.5, space_after=1, left_indent_mm=10, keep_with_next=True)
add_p(doc, "Mã định danh điện tử của tổ chức: 000.03.18.H28", font_size=12.5, space_after=1, left_indent_mm=10, keep_with_next=True)
add_p(doc, "Điện thoại: 0251.3822.456", font_size=12.5, space_after=1, left_indent_mm=10, keep_with_next=True)
add_p(doc, "Email: vanphong@soyte.dongnai.gov.vn", font_size=12.5, space_after=1, left_indent_mm=10, keep_with_next=True)
add_p(doc, "Địa chỉ liên hệ: Số 02 Đường Phan Chu Trinh, Phường Quang Vinh, TP. Biên Hòa, Tỉnh Đồng Nai.", font_size=12.5, space_after=10, left_indent_mm=10, keep_with_next=True)

# Bảng Chữ ký
tbl_sig = doc.add_table(rows=1, cols=2)
tbl_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_sig.autofit = False
c_l = tbl_sig.cell(0, 0)
c_r = tbl_sig.cell(0, 1)
c_l.width, c_r.width = Mm(70), Mm(92)

set_cell_margins(c_r, top=0, bottom=0, left=0, right=0)
p_s1 = c_r.paragraphs[0]
p_s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_s1.paragraph_format.space_before, p_s1.paragraph_format.space_after, p_s1.paragraph_format.line_spacing = Pt(6), Pt(2), 1.15
r = p_s1.add_run("ĐẠI DIỆN TỔ CHỨC ĐẶT HÀNG")
r.font.size, r.font.bold = Pt(13), True

p_s2 = c_r.add_paragraph()
p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_s2.paragraph_format.space_before, p_s2.paragraph_format.space_after, p_s2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = p_s2.add_run("(Họ, tên và chữ ký - đóng dấu)")
r.font.size, r.font.italic = Pt(12.5), True

temp_docx = os.path.join(OUT_DIR, "temp_demo_form1.docx")
final_docx = os.path.join(OUT_DIR, "Mau_I01_DHNV_Phieu_dat_hang_Demo_Trinh_Chieu.docx")
final_pdf = os.path.join(OUT_DIR, "Mau_I01_DHNV_Phieu_dat_hang_Demo_Trinh_Chieu.pdf")

doc.save(temp_docx)
inject_footnotes_into_docx(temp_docx, final_docx, {1: "Phiếu đặt hàng được trình bày không quá 10 trang giấy khổ A4."})
os.remove(temp_docx)
print(f"  [DOCX Saved] {final_docx}")
export_docx_to_pdf(final_docx, final_pdf)

print("\n================================================================================")
print("HOÀN TẤT TẠO BỘ DỮ LIỆU DEMO AN TOÀN TRÌNH CHIẾU VÀO MẪU I.01-ĐHNV!")
print("================================================================================")
