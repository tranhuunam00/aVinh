import os
import sys
import zipfile
import docx
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import win32com.client
import pymupdf
import time
import gc
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

sys.stdout.reconfigure(encoding='utf-8')

BASE_DIR = r"d:\DAOGROUP_WORKSPACE\aVinh"
OUT_DIR = os.path.join(BASE_DIR, "worddata", "chinhthuc_filled")
os.makedirs(OUT_DIR, exist_ok=True)
CREDS_FILE = os.path.join(BASE_DIR, "google-service-account.json")

# Google Drive File IDs
DRIVE_FILES = {
    "form1": {
        "id": "1-kGJVhYaMkPbc2LhbcIKESqnAquyni3M",
        "name": "Mau_I01_DHNV_Phieu_dat_hang.docx"
    },
    "form2": {
        "id": "1Shb5bUAkRvYaFmRlObWb8bSopv0kVLUV",
        "name": "Mau_02_TTMH_To_trinh_phe_duyet.docx"
    },
    "form3": {
        "id": "19m0_7otlDjuTxS1F98CfnbPkDbqlO751",
        "name": "Mau_02_TMMH_Thuyet_minh_nhiem_vu.docx"
    },
    "form4": {
        "id": "1aSMCey4yxB2FtwNspAJPupoGB6lD9DFT",
        "name": "Phu_luc_Giai_trinh_chi_tiet_khoan_chi.docx"
    }
}

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="000000", sz="4"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_p(doc, text="", bold=False, italic=False, font_size=13, align=WD_ALIGN_PARAGRAPH.LEFT, 
          space_before=0, space_after=2, line_spacing=1.15, first_line_indent_mm=0, left_indent_mm=0, keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.keep_with_next = keep_with_next
    if first_line_indent_mm > 0:
        p.paragraph_format.first_line_indent = Mm(first_line_indent_mm)
    if left_indent_mm > 0:
        p.paragraph_format.left_indent = Mm(left_indent_mm)
    if text:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p

def create_base_document(orientation="PORTRAIT", left_mm=28, right_mm=18, top_mm=18, bottom_mm=18):
    doc = docx.Document()
    section = doc.sections[0]
    if orientation == "PORTRAIT":
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT
    else:
        section.page_width = Mm(297)
        section.page_height = Mm(210)
        section.orientation = WD_ORIENT.LANDSCAPE
        
    section.top_margin = Mm(top_mm)
    section.bottom_margin = Mm(bottom_mm)
    section.left_margin = Mm(left_mm)
    section.right_margin = Mm(right_mm)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(13)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(2)
    normal_style.paragraph_format.line_spacing = 1.15
    return doc

def inject_footnotes_into_docx(in_docx_path, out_docx_path, footnotes_dict):
    fn_items = ""
    for fn_id, fn_text in footnotes_dict.items():
        fn_items += f"""
  <w:footnote w:id="{fn_id}">
    <w:p>
      <w:pPr>
        <w:pStyle w:val="FootnoteText"/>
        <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
          <w:rStyle w:val="FootnoteReference"/>
          <w:vertAlign w:val="superscript"/>
          <w:sz w:val="20"/>
        </w:rPr>
        <w:footnoteRef/>
      </w:r>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
          <w:i/>
          <w:sz w:val="22"/>
        </w:rPr>
        <w:t xml:space="preserve"> {fn_text}</w:t>
      </w:r>
    </w:p>
  </w:footnote>"""

    footnotes_xml_content = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:footnote w:type="separator" w:id="-1">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>{fn_items}
</w:footnotes>""".encode('utf-8')

    staging_path = out_docx_path + ".tmp"
    with zipfile.ZipFile(in_docx_path, 'r') as zin:
        with zipfile.ZipFile(staging_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    ct_str = data.decode('utf-8')
                    if 'footnotes.xml' not in ct_str:
                        fn_override = '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                        ct_str = ct_str.replace('</Types>', f'{fn_override}</Types>')
                        data = ct_str.encode('utf-8')
                elif item.filename == 'word/_rels/document.xml.rels':
                    rels_str = data.decode('utf-8')
                    if 'footnotes.xml' not in rels_str:
                        fn_rel = '<Relationship Id="rIdFootnotes" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
                        rels_str = rels_str.replace('</Relationships>', f'{fn_rel}</Relationships>')
                        data = rels_str.encode('utf-8')
                zout.writestr(item, data)
            zout.writestr('word/footnotes.xml', footnotes_xml_content)

    for _ in range(5):
        try:
            if os.path.exists(out_docx_path):
                try:
                    os.remove(out_docx_path)
                except Exception:
                    pass
            os.replace(staging_path, out_docx_path)
            break
        except Exception:
            time.sleep(0.5)

def export_docx_to_pdf(docx_path, pdf_path):
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc_word = None
    try:
        doc_word = word.Documents.Open(os.path.abspath(docx_path))
        doc_word.ExportAsFixedFormat(os.path.abspath(pdf_path), 17)
        doc_word.Close(False)
        doc_word = None
        print(f"  [PDF Exported] {pdf_path}")
    finally:
        word.Quit()
        if doc_word is not None:
            try:
                doc_word.Close(False)
            except Exception:
                pass
            del doc_word
        del word
        gc.collect()
        time.sleep(0.3)

# ==============================================================================
# 1. BIỂU MẪU 1: PHIẾU ĐẶT HÀNG NHIỆM VỤ KH&CN (Mẫu I.01-ĐHNV)
# ==============================================================================
def generate_filled_form1():
    print("\n>>> 1. Đang tạo Mẫu I.01-ĐHNV (Filled)...")
    doc = create_base_document(orientation="PORTRAIT", left_mm=28, right_mm=18, top_mm=18, bottom_mm=18)

    add_p(doc, "PHỤ LỤC II. BIỂU MẪU", bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1, line_spacing=1.1)
    add_p(doc, "Mẫu I.01-ĐHNV", bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=3, line_spacing=1.1)

    top_table = doc.add_table(rows=1, cols=2)
    top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    top_table.autofit = False
    cl, cr = top_table.cell(0, 0), top_table.cell(0, 1)
    cl.width, cr.width = Mm(72), Mm(92)

    set_cell_margins(cl, top=0, bottom=0, left=0, right=0)
    p = cl.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(1), 1.1
    r = p.add_run("SỞ Y TẾ TỈNH HƯNG YÊN\nBỆNH VIỆN ĐA KHOA QUỐC TẾ\nVINMEC OCEAN PARK 2")
    r.font.size, r.font.bold = Pt(11), True

    p_line = cl.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_before, p_line.paragraph_format.space_after, p_line.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
    r_l = p_line.add_run("--------------------")
    r_l.font.size, r_l.font.bold = Pt(10), True

    set_cell_margins(cr, top=0, bottom=0, left=0, right=0)
    p_m1 = cr.paragraphs[0]
    p_m1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_m1.paragraph_format.space_before, p_m1.paragraph_format.space_after, p_m1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.1
    r = p_m1.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r.font.size, r.font.bold = Pt(11), True

    p_m2 = cr.add_paragraph()
    p_m2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_m2.paragraph_format.space_before, p_m2.paragraph_format.space_after, p_m2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.1
    r = p_m2.add_run("Độc lập - Tự do - Hạnh phúc")
    r.font.size, r.font.bold = Pt(12), True

    p_m3 = cr.add_paragraph()
    p_m3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_m3.paragraph_format.space_before, p_m3.paragraph_format.space_after, p_m3.paragraph_format.line_spacing = Pt(0), Pt(2), 1.0
    r = p_m3.add_run("_________________")
    r.font.size, r.font.bold = Pt(10), True

    p_d = cr.add_paragraph()
    p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_d.paragraph_format.space_before, p_d.paragraph_format.space_after, p_d.paragraph_format.line_spacing = Pt(0), Pt(0), 1.1
    r_d2 = p_d.add_run("Hưng Yên, ngày 24 tháng 08 năm 2026")
    r_d2.font.size, r_d2.font.italic = Pt(12), True

    p_title = add_p(doc, "", bold=True, font_size=13.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10, line_spacing=1.15)
    r_t = p_title.add_run("PHIẾU ĐẶT HÀNG NHIỆM VỤ KHOA HỌC VÀ CÔNG NGHỆ CẤP TỈNH")
    r_t.font.size, r_t.font.bold = Pt(13.5), True
    fn_ref = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        f'<w:rStyle w:val="FootnoteReference"/><w:vertAlign w:val="superscript"/><w:b/><w:sz w:val="22"/></w:rPr>'
        f'<w:footnoteReference w:id="1"/></w:r>'
    )
    p_title._p.append(fn_ref)

    # 1. Tên nhiệm vụ
    p1 = add_p(doc, "1. Tên nhiệm vụ: ", bold=True, font_size=12.5, space_after=2)
    r = p1.add_run("Nghiên cứu dịch tễ học và ứng dụng các kỹ thuật công nghệ cao (PRP, phẫu thuật ít xâm lấn MIS có Robot hỗ trợ và phục hồi chức năng thông minh AI) trong chẩn đoán, điều trị thoái hóa khớp gối tại tỉnh Hưng Yên.")
    r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

    # 2. Thuộc lĩnh vực
    add_p(doc, "2. Thuộc lĩnh vực:", bold=True, font_size=12.5, space_after=1.5)
    tbl_f = doc.add_table(rows=4, cols=2)
    tbl_f.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_f.autofit = False
    fields = [
        ("- Khoa học y, dược", True, "- Khoa học tự nhiên", False),
        ("- Khoa học kỹ thuật và công nghệ", False, "- Khoa học nông nghiệp", False),
        ("- Khoa học xã hội & nhân văn", False, "- Công nghệ chiến lược", False),
        ("", False, "", False)
    ]
    for row_idx, (t1, c1, t2, c2) in enumerate(fields):
        cell1, cell2 = tbl_f.cell(row_idx, 0), tbl_f.cell(row_idx, 1)
        cell1.width, cell2.width = Mm(82), Mm(82)
        for cell, txt, is_checked in [(cell1, t1, c1), (cell2, t2, c2)]:
            set_cell_margins(cell, top=0, bottom=0, left=40, right=40)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(1), 1.12
            p.paragraph_format.tab_stops.add_tab_stop(Mm(76), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            if txt:
                r1 = p.add_run(txt + "\t")
                r1.font.size = Pt(12)
                r2 = p.add_run("☒" if is_checked else "☐")
                r2.font.name, r2.font.size = 'Segoe UI Symbol', Pt(12)
                if is_checked: r2.font.bold = True

    # 3. Thuộc loại hình nhiệm vụ
    add_p(doc, "3. Thuộc loại hình nhiệm vụ", bold=True, font_size=12.5, space_before=2, space_after=1.5)
    tbl_t = doc.add_table(rows=2, cols=2)
    tbl_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_t.autofit = False
    types = [
        ("- Nghiên cứu ứng dụng", False, "- Nghiên cứu cơ bản", False),
        ("- Phát triển công nghệ", True, "- Phát triển giải pháp xã hội", False)
    ]
    for row_idx, (t1, c1, t2, c2) in enumerate(types):
        cell1, cell2 = tbl_t.cell(row_idx, 0), tbl_t.cell(row_idx, 1)
        cell1.width, cell2.width = Mm(82), Mm(82)
        for cell, txt, is_checked in [(cell1, t1, c1), (cell2, t2, c2)]:
            set_cell_margins(cell, top=0, bottom=0, left=40, right=40)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(1), 1.12
            p.paragraph_format.tab_stops.add_tab_stop(Mm(76), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            r1 = p.add_run(txt + "\t")
            r1.font.size = Pt(12)
            r2 = p.add_run("☒" if is_checked else "☐")
            r2.font.name, r2.font.size = 'Segoe UI Symbol', Pt(12)
            if is_checked: r2.font.bold = True

    # 4. Căn cứ đặt hàng nhiệm vụ
    add_p(doc, "4. Căn cứ đặt hàng nhiệm vụ:", bold=True, font_size=12.5, space_before=2, space_after=1.5)
    add_p(doc, "- Quyết định số 24/2026/QĐ-UBND ngày 5/6/2026 của UBND tỉnh Hưng Yên quy định quản lý chương trình, nhiệm vụ khoa học, công nghệ và đổi mới sáng tạo sử dụng ngân sách nhà nước;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "- Kế hoạch số 277/KH-UBND ngày 11/7/2026 về phát triển khoa học, công nghệ và đổi mới sáng tạo tỉnh Hưng Yên giai đoạn 2026-2030;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "- Nghị quyết số 57-NQ/TW ngày 22/12/2024 của Bộ Chính trị về phát triển và ứng dụng công nghệ sinh học và y tế kỹ thuật cao phục vụ bảo vệ sức khỏe nhân dân trong tình hình mới;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "- Kế hoạch số 65-Ctr/TU ngày 22/04/2026 của Ban Thường vụ Tỉnh ủy Hưng Yên thực hiện Chương trình phát triển khoa học công nghệ, ứng dụng kỹ thuật cao trong y tế cộng đồng;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "- Chương trình hành động số 01-Ctr/TU ngày 20/10/2025 của Tỉnh ủy Hưng Yên về thúc đẩy nghiên cứu khoa học xã hội, nhân văn và y tế dự phòng giai đoạn 2025-2030.", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=10)

    # 5. Tính cấp thiết
    add_p(doc, "5. Tính cấp thiết của nhiệm vụ:", bold=True, font_size=12.5, space_after=1.5)
    add_p(doc, "Thoái hóa khớp gối (THKG) là một trong những thách thức y tế lớn nhất liên quan đến già hóa dân số và lao động nặng. Tỷ lệ mắc tại Việt Nam ở người trên 40 tuổi chiếm tới hơn 30%, và trên 60% ở người trên 65 tuổi. Tại Hưng Yên, chưa có nghiên cứu dịch tễ học mang tính hệ thống toàn diện tại cộng đồng. Bệnh viện ĐKQT Vinmec Ocean Park 2 với hạ tầng hiện đại, phòng Lab đạt chuẩn tách chiết PRP, hệ thống phòng mổ Robot chấn thương chỉnh hình và trung tâm PHCN tích hợp AI, đề xuất ứng dụng đồng bộ 04 giải pháp công nghệ cao hàng đầu (PRP tự thân, phẫu thuật ít xâm lấn MIS, phẫu thuật Robot định vị 3D và Robot/AI phục hồi chức năng), giúp bảo tồn cấu trúc khớp gối, nâng cao vượt trội chất lượng sống cho người dân ngay tại địa phương.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=12.7)

    # 6. Mục tiêu
    add_p(doc, "6. Mục tiêu:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
    add_p(doc, "- Mục tiêu chung: Đánh giá toàn diện thực trạng dịch tễ học, đặc điểm lâm sàng, cận lâm sàng, yếu tố liên quan của bệnh lý THKG tại cộng đồng tỉnh Hưng Yên; đồng thời ứng dụng chuỗi giải pháp công nghệ cao đồng bộ (PRP tự thân, MIS, Robot hỗ trợ, Robot & AI PHCN) nhằm tối ưu hóa chất lượng điều trị, giảm tỷ lệ tàn phế cho người bệnh.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "- Mục tiêu cụ thể: (1) Xác định tỷ lệ hiện mắc và các yếu tố liên quan của THKG ở người từ 40 tuổi trở lên tại Hưng Yên; (2) Mô tả đặc điểm lâm sàng và hình ảnh học (X-quang Kellgren-Lawrence, MRI, siêu âm); (3) Ứng dụng điều trị công nghệ cao (PRP cho giai đoạn I-III, MIS và Robot cho giai đoạn III-IV); (4) Đánh giá hiệu quả phục hồi chức năng của Robot thông minh & AI; (5) Xây dựng bộ khuyến nghị khoa học và chuyển giao kỹ thuật cho mạng lưới y tế tỉnh Hưng Yên.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=10)

    # 7. Dự kiến các nội dung chính cần thực hiện
    add_p(doc, "7. Dự kiến các nội dung chính cần thực hiện:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
    add_p(doc, "• Nội dung 1: Xây dựng đề cương, công cụ khảo sát (WOMAC, KOOS, VAS) và chuẩn hóa quy trình kỹ thuật lâm sàng công nghệ cao, thông qua Hội đồng Đạo đức;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "• Nội dung 2: Điều tra khám sàng lọc thực địa cộng đồng cho 1.500 - 2.000 người dân từ 40 tuổi trở lên tại 10 huyện/thành phố thuộc tỉnh Hưng Yên;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "• Nội dung 3: Khám chuyên khoa sâu và chẩn đoán xác định cận lâm sàng kỹ thuật cao (X-quang thẳng/nghiêng chịu lực, MRI, siêu âm đo sụn) tại BV ĐKQT Vinmec Ocean Park 2;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "• Nội dung 4: Phân tích số liệu dịch tễ học, tính tỷ lệ hiện mắc và phân tích hồi quy đa biến Logistic các yếu tố liên quan (BMI, nghề nghiệp, thói quen ngồi xổm, bệnh nền);", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "• Nội dung 5: Triển khai can thiệp điều trị chuỗi công nghệ cao (50-80 ca PRP tự thân, 30 ca MIS, 20-30 ca phẫu thuật Robot, chương trình Robot/AI PHCN) và theo dõi dọc 1, 3, 6, 12 tháng;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "• Nội dung 6: Xây dựng tài liệu kỹ thuật khuyến nghị, tổ chức hội thảo khoa học tỉnh và tập huấn chuyển giao quy trình điều trị bảo tồn PRP cho các BV tuyến tỉnh, huyện.", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=10)

    # 8. Dự kiến các kết quả thực hiện nhiệm vụ và các chỉ tiêu cần đạt
    add_p(doc, "8. Dự kiến các kết quả thực hiện nhiệm vụ và các chỉ tiêu cần đạt:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
    add_p(doc, "1. 01 Đề cương hoàn chỉnh và bộ công cụ khảo sát dịch tễ học (WOMAC, KOOS, VAS);", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "2. Cơ sở dữ liệu dịch tễ số hóa của ít nhất 1.500 người dân từ 40 tuổi trở lên tại Hưng Yên;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "3. 01 Báo cáo khoa học phân tích sâu đặc điểm lâm sàng, cận lâm sàng (X-quang, MRI, siêu âm);", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "4. 04 Quy trình kỹ thuật chuẩn hóa (Quy trình tách chiết & tiêm PRP; Quy trình MIS; Quy trình Robot phẫu thuật; Quy trình Robot/AI phục hồi chức năng);", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "5. Ít nhất 50 bệnh nhân điều trị PRP đạt kết quả tốt (giảm VAS >= 40%, tăng WOMAC >= 30%);", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "6. Ít nhất 30 bệnh nhân mổ MIS và 20 bệnh nhân mổ Robot đạt độ chính xác cao, hồi phục nhanh;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "7. 01 Cuốn tài liệu khuyến nghị kỹ thuật; 02 lớp tập huấn cho 40 bác sĩ y tế cơ sở Hưng Yên;", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "8. Ít nhất 02 bài báo đăng trên tạp chí y học chuyên ngành trong nước và 01 bài báo quốc tế (ISI/Scopus); đào tạo 02 tiến sĩ y khoa.", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=10)

    # 9. Hiệu quả và tác động
    add_p(doc, "9. Hiệu quả và tác động của kết quả thực hiện nhiệm vụ:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
    add_p(doc, "• Hiệu quả KH&CN: Cung cấp số liệu dịch tễ học thực chứng đầu tiên tại Hưng Yên, chuẩn hóa chuỗi quy trình công nghệ cao phục vụ chăm sóc sức khỏe lâu dài;\n• Hiệu quả KT-XH: Giảm ngày nằm viện từ 7-10 ngày xuống 3-5 ngày nhờ MIS/Robot, hạn chế chi phí thuốc giảm đau và mổ lại (tiết kiệm hàng trăm triệu đồng/ca), giúp bệnh nhân sớm phục hồi lao động;\n• Nâng cao năng lực y tế: BV ĐKQT Vinmec Ocean Park 2 là hạt nhân đào tạo thực hành chuyển giao kỹ thuật cho y tế tuyến dưới tỉnh Hưng Yên.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3, first_line_indent_mm=12.7)

    # 10. Dự kiến kinh phí thực hiện
    p10 = add_p(doc, "10. Dự kiến kinh phí thực hiện: ", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
    r = p10.add_run("18.425.000.000 VNĐ (Mười tám tỷ bốn trăm hai mươi lăm triệu đồng chẵn).")
    r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

    # 11. Dự kiến thời gian thực hiện
    p11 = add_p(doc, "11. Dự kiến thời gian thực hiện: ", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
    r = p11.add_run("24 tháng (Từ tháng 11 năm 2026 đến hết tháng 10 năm 2028).")
    r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

    # Sub-bullets of kinh phí
    p_ns = add_p(doc, "- Kinh phí hỗ trợ từ NSNN: ", bold=True, font_size=12.5, space_after=1.5, left_indent_mm=8, keep_with_next=True)
    r = p_ns.add_run("4.500.000.000 VNĐ (Bốn tỷ năm trăm triệu đồng chẵn - 24.4%).")
    r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

    p_nk = add_p(doc, "- Kinh phí từ nguồn khác: ", bold=True, font_size=12.5, space_after=2.5, left_indent_mm=8, keep_with_next=True)
    r = p_nk.add_run("13.925.000.000 VNĐ (Nguồn đối ứng tự có của BV ĐKQT Vinmec Ocean Park 2 / Hệ thống Y tế Vinmec - 75.6%).")
    r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

    # 12. Đề xuất tổ chức được xét giao trực tiếp
    p12 = add_p(doc, "12. Đề xuất tổ chức được xét giao trực tiếp (nếu có): ", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
    r = p12.add_run("Bệnh viện Đa khoa Quốc tế Vinmec Ocean Park 2 (Địa chỉ: Xã Nghĩa Trụ, Tỉnh Hưng Yên).")
    r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(12.5), False

    # 13. Tiếp nhận và phương án quản lý
    add_p(doc, "13. Tiếp nhận và phương án tổ chức quản lý, sử dụng kết quả của nhiệm vụ của cơ quan đặt hàng:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
    add_p(doc, "(Nêu rõ việc Nhà nước sẽ nắm giữ quyền quản lý, sử dụng, quyền sở hữu kết quả trong trường hợp nhà nước có yêu cầu tiếp nhận kết quả)", italic=True, font_size=11.5, space_after=1.5, left_indent_mm=5)
    add_p(doc, "- Cơ quan đặt hàng tiếp nhận kết quả: Sở Y tế tỉnh Hưng Yên;\n- Sở Y tế trực tiếp quản lý, lưu trữ báo cáo dịch tễ học định hướng phát triển mạng lưới cơ xương khớp tỉnh giai đoạn 2026-2030;\n- Tiếp nhận bộ tài liệu kỹ thuật khuyến nghị để nhân rộng áp dụng như hướng dẫn chuyên môn chính thức;\n- Phối hợp với Vinmec Ocean Park 2 chuyển giao kỹ thuật tiêm PRP tự thân cho ít nhất 3-5 bệnh viện đa khoa tuyến tỉnh, huyện;\n- Phối hợp Sở KH&CN Hưng Yên số hóa bộ dữ liệu đưa vào kho dữ liệu KH&CN dùng chung.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2.5, first_line_indent_mm=10)

    # 14. Danh mục tài liệu tham khảo
    add_p(doc, "14. Danh mục tài liệu tham khảo:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
    add_p(doc, "[1] Bộ Y tế Việt Nam (2020), Hướng dẫn chẩn đoán và điều trị các bệnh về cơ xương khớp, NXB Y học.\n[2] Bệnh viện ĐKQT Vinmec (2024), Hướng dẫn lâm sàng ứng dụng PRP tự thân trong điều trị THKG giai đoạn I-III.\n[3] UBND tỉnh Hưng Yên (2026), Quyết định số 24/2026/QĐ-UBND ngày 5/6/2026.\n[4] UBND tỉnh Hưng Yên (2026), Kế hoạch số 277/KH-UBND ngày 11/7/2026.\n[5] Bannuru R. R. et al. (2019), OARSI guidelines for knee osteoarthritis, Osteoarthritis Cartilage, 27(11).\n[6] Kolasinski S. L. et al. (2020), ACR/AF Guideline for Management of Osteoarthritis, Arthritis Care Res, 72(2).\n[7] Filardo G. et al. (2021), PRP intra-articular injections for osteoarthritis, Am J Sports Med, 49(2).\n[8] Kayani B. et al. (2020), Robotic-arm assisted total knee arthroplasty, Bone Joint J, 102-B(11).\n[9] Chen J. et al. (2022), AI and robotic rehabilitation in post-operative knee arthroplasty, J Orthop Surg Res, 17(1).\n[10] Nguyễn Văn Hùng, Trần Trung Dũng (2021), Thay khớp gối toàn phần lối mổ ít xâm lấn bảo tồn cơ tứ đầu đùi, Tạp chí Y học VN.", 
          font_size=11.5, space_after=2.5, first_line_indent_mm=10)

    # 15. Thông tin liên hệ
    add_p(doc, "15. Thông tin liên hệ:", bold=True, font_size=12.5, space_after=1.5, keep_with_next=True)
    add_p(doc, "Tên tổ chức: SỞ Y TẾ TỈNH HƯNG YÊN", font_size=12.5, space_after=1, left_indent_mm=10, keep_with_next=True)
    add_p(doc, "Đại diện tổ chức: Lãnh đạo Sở Y tế tỉnh Hưng Yên", font_size=12.5, space_after=1, left_indent_mm=10, keep_with_next=True)
    add_p(doc, "Địa chỉ liên hệ: Tỉnh Hưng Yên.", font_size=12.5, space_after=10, left_indent_mm=10, keep_with_next=True)

    # Bảng Chữ ký
    tbl_sig = doc.add_table(rows=1, cols=2)
    tbl_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_sig.autofit = False
    c_l, c_r = tbl_sig.cell(0, 0), tbl_sig.cell(0, 1)
    c_l.width, c_r.width = Mm(70), Mm(92)
    set_cell_margins(c_r, top=0, bottom=0, left=0, right=0)
    p_s1 = c_r.paragraphs[0]
    p_s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s1.paragraph_format.space_before, p_s1.paragraph_format.space_after, p_s1.paragraph_format.line_spacing = Pt(6), Pt(2), 1.15
    r = p_s1.add_run("ĐẠI DIỆN TỔ CHỨC ĐẶT HÀNG")
    r.font.size, r.font.bold = Pt(13), True
    p_s2 = c_r.add_paragraph()
    p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s2.paragraph_format.space_before, p_s2.paragraph_format.space_after, p_s2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
    r = p_s2.add_run("(Họ, tên và chữ ký - đóng dấu)")
    r.font.size, r.font.italic = Pt(12.5), True

    temp_docx = os.path.join(OUT_DIR, "temp_form1.docx")
    final_docx = os.path.join(OUT_DIR, "Mau_I01_DHNV_Phieu_dat_hang.docx")
    final_pdf = os.path.join(OUT_DIR, "Mau_I01_DHNV_Phieu_dat_hang.pdf")
    doc.save(temp_docx)
    inject_footnotes_into_docx(temp_docx, final_docx, {1: "Phiếu đặt hàng được trình bày không quá 10 trang giấy khổ A4."})
    if os.path.exists(temp_docx): os.remove(temp_docx)
    print(f"  [DOCX Saved] {final_docx}")
    export_docx_to_pdf(final_docx, final_pdf)
    return final_docx

# ==============================================================================
# 2. BIỂU MẪU 2: TỜ TRÌNH PHÊ DUYỆT (Mẫu 02-TTMH)
# ==============================================================================
def generate_filled_form2():
    print("\n>>> 2. Đang tạo Mẫu 02-TTMH (Filled)...")
    doc = create_base_document(orientation="PORTRAIT", left_mm=28, right_mm=18, top_mm=18, bottom_mm=18)

    add_p(doc, "Mẫu 02 -TTMH", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=4, line_spacing=1.15)

    top_table = doc.add_table(rows=1, cols=2)
    top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    top_table.autofit = False
    cl, cr = top_table.cell(0, 0), top_table.cell(0, 1)
    cl.width, cr.width = Mm(65), Mm(99)

    set_cell_margins(cl, top=0, bottom=0, left=0, right=0)
    p_org1 = cl.paragraphs[0]
    p_org1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_org1.paragraph_format.space_before, p_org1.paragraph_format.space_after, p_org1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
    r = p_org1.add_run("SỞ Y TẾ TỈNH HƯNG YÊN")
    r.font.size, r.font.bold = Pt(12), True

    p_org3 = cl.add_paragraph()
    p_org3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_org3.paragraph_format.space_before, p_org3.paragraph_format.space_after, p_org3.paragraph_format.line_spacing = Pt(0), Pt(3), 1.0
    r = p_org3.add_run("-------")
    r.font.size, r.font.bold = Pt(11), True

    p_so = cl.add_paragraph()
    p_so.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_so.paragraph_format.space_before, p_so.paragraph_format.space_after, p_so.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
    r = p_so.add_run("Số: ….../TTr-SYT")
    r.font.size = Pt(12)

    set_cell_margins(cr, top=0, bottom=0, left=0, right=0)
    p_m1 = cr.paragraphs[0]
    p_m1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_m1.paragraph_format.space_before, p_m1.paragraph_format.space_after, p_m1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
    r = p_m1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r.font.size, r.font.bold = Pt(12), True

    p_m2 = cr.add_paragraph()
    p_m2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_m2.paragraph_format.space_before, p_m2.paragraph_format.space_after, p_m2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
    r = p_m2.add_run("Độc lập - Tự do - Hạnh phúc")
    r.font.size, r.font.bold = Pt(13), True

    p_d = cr.add_paragraph()
    p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_d.paragraph_format.space_before, p_d.paragraph_format.space_after, p_d.paragraph_format.line_spacing = Pt(3), Pt(0), 1.15
    r_d = p_d.add_run("Hưng Yên, ngày …. tháng… năm 2026")
    r_d.font.size, r_d.font.italic = Pt(13), True

    # Title
    add_p(doc, "TỜ TRÌNH", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=3, line_spacing=1.15)
    add_p(doc, "Về việc phê duyệt, hỗ trợ nhiệm vụ ứng dụng, chuyển giao tiến bộ", bold=True, font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1, line_spacing=1.15)
    add_p(doc, "khoa học và công nghệ vào đời sống và sản xuất", bold=True, font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0, line_spacing=1.15)
    add_p(doc, "___________________________", bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=10, line_spacing=1.0)

    # Kính gửi
    add_p(doc, "Kính gửi:", bold=False, font_size=13, space_before=0, space_after=2, line_spacing=1.15, left_indent_mm=34)
    add_p(doc, "- Sở Khoa học và Công nghệ;", font_size=13, space_after=2, line_spacing=1.15, left_indent_mm=56)
    add_p(doc, "- Sở Tài chính.", font_size=13, space_after=6, line_spacing=1.15, left_indent_mm=56)

    # Căn cứ
    add_p(doc, "Căn cứ Luật Ngân sách nhà nước ngày 25 tháng 6 năm 2025;", italic=True, font_size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, line_spacing=1.15, first_line_indent_mm=12.7)
    add_p(doc, "Căn cứ Luật Quản lý, sử dụng tài sản công ngày 21 tháng 6 năm 2017;", italic=True, font_size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, line_spacing=1.15, first_line_indent_mm=12.7)
    add_p(doc, "Căn cứ Nghị định số 104/2026/NĐ-CP ngày 31 tháng 3 năm 2026 của Chính phủ quy định việc lập dự toán, quản lý, sử dụng và quyết toán chi thường xuyên để thực hiện các nhiệm vụ quy định tại Điều 40 Luật Ngân sách nhà nước;", italic=True, font_size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, line_spacing=1.15, first_line_indent_mm=12.7)
    add_p(doc, "Căn cứ Quyết định số 24/2026/QĐ-UBND ngày 5/6/2026 và Kế hoạch số 277/KH-UBND ngày 11/7/2026 của UBND tỉnh Hưng Yên về quản lý và phát triển khoa học, công nghệ và đổi mới sáng tạo;", italic=True, font_size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, line_spacing=1.15, first_line_indent_mm=12.7)

    # Đoạn dẫn
    p_intro = add_p(doc, "", font_size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, line_spacing=1.15, first_line_indent_mm=12.7)
    r_i1 = p_intro.add_run("Sở Y tế tỉnh Hưng Yên")
    r_i1.font.italic, r_i1.font.bold = False, True
    r_i2 = p_intro.add_run(" kính trình Sở Khoa học và Công nghệ, Sở Tài chính xem xét thẩm định, phê duyệt hỗ trợ nhiệm vụ ứng dụng, chuyển giao tiến bộ khoa học và công nghệ vào đời sống và sản xuất năm 2027 cho ")
    r_i3 = p_intro.add_run("Bệnh viện Đa khoa Quốc tế Vinmec Ocean Park 2")
    r_i3.font.italic, r_i3.font.bold = False, True
    r_i4 = p_intro.add_run(" với các nội dung chủ yếu như sau:")

    # I. THÔNG TIN CHUNG
    add_p(doc, "I. THÔNG TIN CHUNG", bold=True, font_size=13, space_before=4, space_after=2)
    add_p(doc, "1. Tên cơ quan, đơn vị đề xuất: Bệnh viện Đa khoa Quốc tế Vinmec Ocean Park 2 (phối hợp Sở Y tế tỉnh Hưng Yên).", font_size=13, space_after=2, first_line_indent_mm=10)
    add_p(doc, "2. Tên nhiệm vụ: Nghiên cứu dịch tễ học và ứng dụng các kỹ thuật công nghệ cao (PRP, phẫu thuật ít xâm lấn MIS có Robot hỗ trợ và phục hồi chức năng thông minh AI) trong chẩn đoán, điều trị thoái hóa khớp gối tại tỉnh Hưng Yên.", font_size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=10)
    add_p(doc, "3. Dự kiến kinh phí: 18.425 triệu đồng. Trong đó:", font_size=13, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "- Nguồn ngân sách nhà nước: 4.500 triệu đồng", font_size=13, space_after=1.5, left_indent_mm=18)
    add_p(doc, "- Nguồn khác: 13.925 triệu đồng (Vốn đối ứng tự có của đơn vị thực hiện)", font_size=13, space_after=2, left_indent_mm=18)

    p_tg = add_p(doc, "4. Thời gian thực hiện nhiệm vụ: ", font_size=13, space_after=2, first_line_indent_mm=10)
    r_tg = p_tg.add_run("24 tháng (Từ tháng 11 năm 2026 đến tháng 10 năm 2028).")
    fn_ref2 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        f'<w:rStyle w:val="FootnoteReference"/><w:vertAlign w:val="superscript"/><w:sz w:val="20"/></w:rPr>'
        f'<w:footnoteReference w:id="2"/></w:r>'
    )
    p_tg._p.append(fn_ref2)

    add_p(doc, "5. Các nội dung khác: Ứng dụng chuỗi công nghệ cao PRP tự thân, MIS, phẫu thuật Robot và Robot/AI phục hồi chức năng; đào tạo chuyển giao kỹ thuật cho các BV đa khoa tuyến tỉnh và huyện.", font_size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=10)

    # II. NGUỒN KINH PHÍ THỰC HIỆN
    add_p(doc, "II. NGUỒN KINH PHÍ THỰC HIỆN", bold=True, font_size=13, space_before=4, space_after=2)
    add_p(doc, "Đề nghị hỗ trợ từ nguồn ngân sách sự nghiệp KHCN, ĐMST&CĐS năm 2027 của tỉnh Hưng Yên: 4.500.000.000 đồng; Vốn đối ứng của Hệ thống Y tế Vinmec: 13.925.000.000 đồng.", font_size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=12.7)

    # III. HỒ SƠ TÀI LIỆU KÈM THEO
    add_p(doc, "III. HỒ SƠ TÀI LIỆU KÈM THEO", bold=True, font_size=13, space_before=4, space_after=2)
    add_p(doc, "- Thuyết minh chi tiết nhiệm vụ kèm theo dự toán;", font_size=13, space_after=2, left_indent_mm=10)
    add_p(doc, "- Báo giá và các tài liệu liên quan khác kèm theo.", font_size=13, space_after=4, left_indent_mm=10)

    # Đoạn kết
    p_end = add_p(doc, "", font_size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10, line_spacing=1.15, first_line_indent_mm=12.7)
    r_e1 = p_end.add_run("Sở Y tế tỉnh Hưng Yên")
    r_e1.font.italic, r_e1.font.bold = False, True
    r_e2 = p_end.add_run(" kính trình Sở Khoa học và Công nghệ, Sở Tài chính xem xét thẩm định, trình UBND tỉnh phê duyệt./.")

    # Bảng Chữ ký & Nơi nhận
    tbl_bot = doc.add_table(rows=1, cols=2)
    tbl_bot.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_bot.autofit = False
    c_noi_nhan, c_cq = tbl_bot.cell(0, 0), tbl_bot.cell(0, 1)
    c_noi_nhan.width, c_cq.width = Mm(75), Mm(89)

    set_cell_margins(c_noi_nhan, top=0, bottom=0, left=0, right=0)
    p_nn1 = c_noi_nhan.paragraphs[0]
    p_nn1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_nn1.paragraph_format.space_before, p_nn1.paragraph_format.space_after, p_nn1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
    r = p_nn1.add_run("Nơi nhận:")
    r.font.size, r.font.bold, r.font.italic = Pt(12), True, True

    for text_nn in ["- Như trên;", "- UBND tỉnh Hưng Yên (để b/c);", "- Lưu: VT, KHTC."]:
        p = c_noi_nhan.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
        r = p.add_run(text_nn)
        r.font.size = Pt(11)

    set_cell_margins(c_cq, top=0, bottom=0, left=0, right=0)
    p_cq1 = c_cq.paragraphs[0]
    p_cq1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cq1.paragraph_format.space_before, p_cq1.paragraph_format.space_after, p_cq1.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
    r = p_cq1.add_run("GIÁM ĐỐC SỞ Y TẾ TỈNH HƯNG YÊN")
    r.font.size, r.font.bold = Pt(13), True

    p_cq2 = c_cq.add_paragraph()
    p_cq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cq2.paragraph_format.space_before, p_cq2.paragraph_format.space_after, p_cq2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
    r = p_cq2.add_run("(Ký, ghi rõ họ tên, chức vụ và đóng dấu)")
    r.font.size, r.font.italic = Pt(12), True

    temp_docx2 = os.path.join(OUT_DIR, "temp_form2.docx")
    final_docx2 = os.path.join(OUT_DIR, "Mau_02_TTMH_To_trinh_phe_duyet.docx")
    final_pdf2 = os.path.join(OUT_DIR, "Mau_02_TTMH_To_trinh_phe_duyet.pdf")
    doc.save(temp_docx2)
    inject_footnotes_into_docx(temp_docx2, final_docx2, {
        2: "Trong trường hợp nhiệm vụ có thời gian thực hiện trên 01 năm thì cần xác định dự toán kinh phí thực hiện trong từng năm."
    })
    if os.path.exists(temp_docx2): os.remove(temp_docx2)
    print(f"  [DOCX Saved] {final_docx2}")
    export_docx_to_pdf(final_docx2, final_pdf2)
    return final_docx2

# ==============================================================================
# 3. BIỂU MẪU 3: THUYẾT MINH NHIỆM VỤ (Mẫu 02-TMMH)
# ==============================================================================
def generate_filled_form3():
    print("\n>>> 3. Đang tạo Mẫu 02-TMMH (Filled)...")
    doc = create_base_document(orientation="PORTRAIT", left_mm=28, right_mm=18, top_mm=18, bottom_mm=18)

    add_p(doc, "Mẫu 02 -TMMH", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=4, line_spacing=1.15)

    top_table = doc.add_table(rows=1, cols=2)
    top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    top_table.autofit = False
    cl, cr = top_table.cell(0, 0), top_table.cell(0, 1)
    cl.width, cr.width = Mm(72), Mm(92)

    set_cell_margins(cl, top=0, bottom=0, left=0, right=0)
    p_org1 = cl.paragraphs[0]
    p_org1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_org1.paragraph_format.space_before, p_org1.paragraph_format.space_after, p_org1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
    r = p_org1.add_run("SỞ Y TẾ TỈNH HƯNG YÊN\nBỆNH VIỆN ĐA KHOA QUỐC TẾ\nVINMEC OCEAN PARK 2")
    r.font.size, r.font.bold = Pt(11), True

    p_org3 = cl.add_paragraph()
    p_org3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_org3.paragraph_format.space_before, p_org3.paragraph_format.space_after, p_org3.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
    r = p_org3.add_run("--------------------")
    r.font.size, r.font.bold = Pt(10), True

    set_cell_margins(cr, top=0, bottom=0, left=0, right=0)
    p_m1 = cr.paragraphs[0]
    p_m1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_m1.paragraph_format.space_before, p_m1.paragraph_format.space_after, p_m1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
    r = p_m1.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r.font.size, r.font.bold = Pt(11), True

    p_m2 = cr.add_paragraph()
    p_m2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_m2.paragraph_format.space_before, p_m2.paragraph_format.space_after, p_m2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
    r = p_m2.add_run("Độc lập - Tự do - Hạnh phúc")
    r.font.size, r.font.bold = Pt(12), True

    p_m3 = cr.add_paragraph()
    p_m3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_m3.paragraph_format.space_before, p_m3.paragraph_format.space_after, p_m3.paragraph_format.line_spacing = Pt(0), Pt(2), 1.0
    r = p_m3.add_run("_________________")
    r.font.size, r.font.bold = Pt(10), True

    p_d = cr.add_paragraph()
    p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_d.paragraph_format.space_before, p_d.paragraph_format.space_after, p_d.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
    r_d2 = p_d.add_run("Hưng Yên, ngày 24 tháng 08 năm 2026")
    r_d2.font.size, r_d2.font.italic = Pt(12), True

    add_p(doc, "THUYẾT MINH NHIỆM VỤ ỨNG DỤNG, NHÂN RỘNG KẾT QUẢ NGHIÊN CỨU KHOA HỌC VÀ CÔNG NGHỆ NĂM 2027", 
          bold=True, font_size=13.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=8, line_spacing=1.15)

    # I. THÔNG TIN CHUNG VỀ NHIỆM VỤ
    add_p(doc, "I. THÔNG TIN CHUNG VỀ NHIỆM VỤ", bold=True, font_size=13, space_before=4, space_after=2)
    add_p(doc, "1. Tên nhiệm vụ: Nghiên cứu dịch tễ học và ứng dụng các kỹ thuật công nghệ cao (PRP, phẫu thuật ít xâm lấn MIS có Robot hỗ trợ và phục hồi chức năng thông minh AI) trong chẩn đoán, điều trị thoái hóa khớp gối tại tỉnh Hưng Yên.", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "2. Địa điểm thực hiện: Tỉnh Hưng Yên (Khảo sát dịch tễ 10 huyện/TP; Can thiệp kỹ thuật cao tại BV ĐKQT Vinmec Ocean Park 2; Chuyển giao tại các BVĐK tuyến huyện).", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "3. Thời gian thực hiện: 24 tháng (Từ tháng 11/2026 đến tháng 10/2028).", font_size=12.5, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "4. Cơ quan quản lý: Sở Y tế tỉnh Hưng Yên.", font_size=12.5, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "5. Cơ quan chủ trì / Đơn vị thực hiện: Bệnh viện Đa khoa Quốc tế Vinmec Ocean Park 2.\n- Đại diện: Ban Giám đốc Bệnh viện Đa khoa Quốc tế Vinmec Ocean Park 2;\n- Địa chỉ: Xã Nghĩa Trụ, Huyện Văn Giang, Tỉnh Hưng Yên.", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "6. Tổng kinh phí thực hiện: 18.425 triệu đồng (Mười tám tỷ bốn trăm hai mươi lăm triệu đồng chẵn).\n- Kinh phí hỗ trợ từ NS Sự nghiệp KH&CN: 4.500 triệu đồng (24.4%);\n- Kinh phí đối ứng tự có của đơn vị: 13.925 triệu đồng (75.6%).", font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3, first_line_indent_mm=10)

    # II. NỘI DUNG VÀ TIẾN ĐỘ THỰC HIỆN
    add_p(doc, "II. NỘI DUNG VÀ TIẾN ĐỘ THỰC HIỆN", bold=True, font_size=13, space_before=4, space_after=2)
    add_p(doc, "1. Căn cứ pháp lý, tính cấp thiết và cơ sở khoa học:", bold=True, font_size=12.5, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "Thoái hóa khớp gối là nguyên nhân hàng đầu gây hạn chế vận động và tàn phế ở người trung niên, cao tuổi. Tại Hưng Yên, chưa có điều tra dịch tễ học diện rộng để xác định tỷ lệ mắc và các yếu tố nguy cơ. Việc ứng dụng đồng bộ 04 công nghệ y học hàng đầu thế giới (liệu pháp sinh học PRP tự thân, phẫu thuật tối thiểu MIS, phẫu thuật Robot cắt xương 3D và Robot/AI phục hồi chức năng) sẽ giúp giải quyết triệt để vấn đề từ bảo tồn đến can thiệp chuyên sâu, nâng cao vượt bậc chất lượng điều trị ngay tại tỉnh.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=10)
    add_p(doc, "2. Mục tiêu nhiệm vụ:", bold=True, font_size=12.5, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "Đánh giá dịch tễ học THKG tại Hưng Yên (n=1.500-2.000); ứng dụng điều trị công nghệ cao (50-80 ca PRP, 30 ca MIS, 20-30 ca phẫu thuật Robot, chương trình tập Robot/AI); chuẩn hóa quy trình và chuyển giao kỹ thuật cho y tế tuyến dưới.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=10)
    add_p(doc, "3. Nội dung và phương pháp triển khai:", bold=True, font_size=12.5, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "• Nội dung 1: Xây dựng đề cương, công cụ khảo sát (WOMAC, KOOS, VAS) và chuẩn hóa quy trình lâm sàng;\n• Nội dung 2: Khám sàng lọc cộng đồng cho 1.500 - 2.000 người dân từ 40 tuổi trở lên tại 10 huyện/TP;\n• Nội dung 3: Khám chẩn đoán cận lâm sàng chuyên sâu (X-quang số hóa chịu lực, MRI, siêu âm khớp) tại Vinmec Ocean Park 2;\n• Nội dung 4: Phân tích số liệu dịch tễ học bằng SPSS/Stata, phân tích hồi quy đa biến Logistic;\n• Nội dung 5: Can thiệp điều trị chuỗi công nghệ cao (PRP, MIS, Robot) kết hợp phục hồi chức năng Robot & AI, theo dõi dọc 1, 3, 6, 12 tháng;\n• Nội dung 6: Xây dựng bộ khuyến nghị kỹ thuật, tổ chức hội thảo khoa học và đào tạo chuyển giao cho y tế cơ sở.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=10)
    add_p(doc, "4. Tiến độ thực hiện (24 tháng):", bold=True, font_size=12.5, space_after=1.5, first_line_indent_mm=10)
    add_p(doc, "- Giai đoạn 1 (11/2026 - 03/2027): Xây dựng đề cương, công cụ, thông qua HĐ Đạo đức, tập huấn điều tra viên;\n- Giai đoạn 2 (04/2027 - 09/2027): Điều tra sàng lọc cộng đồng và chẩn đoán cận lâm sàng chuyên sâu;\n- Giai đoạn 3 (10/2027 - 04/2028): Nhập liệu phân tích dịch tễ, can thiệp lâm sàng chuỗi công nghệ cao và tập PHCN AI;\n- Giai đoạn 4 (05/2028 - 10/2028): Theo dõi dọc kết quả, hoàn thiện khuyến nghị, đào tạo chuyển giao và nghiệm thu.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3, first_line_indent_mm=10)

    # III. DỰ KIẾN SẢN PHẨM, PHÂN TÍCH HIỆU QUẢ, KẾ HOẠCH NHÂN RỘNG
    add_p(doc, "III. DỰ KIẾN SẢN PHẨM, PHÂN TÍCH HIỆU QUẢ, KẾ HOẠCH NHÂN RỘNG", bold=True, font_size=13, space_before=4, space_after=2)
    add_p(doc, "1. Dự kiến sản phẩm của nhiệm vụ: 01 Đề cương & bộ công cụ; Cơ sở dữ liệu dịch tễ 1.500 người; 01 Báo cáo đặc điểm lâm sàng cận lâm sàng; 04 Quy trình kỹ thuật công nghệ cao; 50-80 ca PRP, 30 ca MIS, 20 ca phẫu thuật Robot đạt kết quả tốt; 01 Cuốn sách hướng dẫn kỹ thuật; 02 lớp tập huấn 40 bác sĩ; 02 bài báo trong nước, 01 bài báo quốc tế (ISI/Scopus); đào tạo 02 tiến sĩ.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=10)
    add_p(doc, "2. Phân tích hiệu quả: Rút ngắn thời gian nằm viện xuống 3-5 ngày, tăng độ bền khớp nhân tạo trên 20 năm, giảm nguy cơ phẫu thuật lại tiết kiệm hàng trăm triệu đồng/ca, phục hồi khả năng lao động và nâng cao sức khỏe cộng đồng tỉnh Hưng Yên.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=10)
    add_p(doc, "3. Kế hoạch triển khai nhân rộng nhiệm vụ: Sở Y tế ban hành tài liệu khuyến nghị áp dụng toàn tỉnh, Vinmec Ocean Park 2 trực tiếp chuyển giao kỹ thuật tiêm PRP cho 3-5 BVĐK huyện, số hóa dữ liệu đưa vào kho KH&CN tỉnh.", 
          font_size=12.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, first_line_indent_mm=10)

    # Chữ ký 2 bên
    add_p(doc, "Hưng Yên, ngày 24 tháng 08 năm 2026", italic=True, font_size=12.5, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=4, space_after=6)
    tbl_sig3 = doc.add_table(rows=1, cols=2)
    tbl_sig3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_sig3.autofit = False
    cth, cql = tbl_sig3.cell(0, 0), tbl_sig3.cell(0, 1)
    cth.width, cql.width = Mm(81), Mm(81)

    set_cell_margins(cth, top=0, bottom=0, left=0, right=0)
    pth1 = cth.paragraphs[0]
    pth1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pth1.paragraph_format.space_before, pth1.paragraph_format.space_after, pth1.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
    r = pth1.add_run("ĐƠN VỊ THỰC HIỆN")
    r.font.size, r.font.bold = Pt(13), True
    pth2 = cth.add_paragraph()
    pth2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pth2.paragraph_format.space_before, pth2.paragraph_format.space_after, pth2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
    r = pth2.add_run("(Ký và đóng dấu)")
    r.font.size, r.font.italic = Pt(12), True

    set_cell_margins(cql, top=0, bottom=0, left=0, right=0)
    pql1 = cql.paragraphs[0]
    pql1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pql1.paragraph_format.space_before, pql1.paragraph_format.space_after, pql1.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
    r = pql1.add_run("ĐƠN VỊ QUẢN LÝ")
    r.font.size, r.font.bold = Pt(13), True
    pql2 = cql.add_paragraph()
    pql2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pql2.paragraph_format.space_before, pql2.paragraph_format.space_after, pql2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
    r = pql2.add_run("(Ký và đóng dấu)")
    r.font.size, r.font.italic = Pt(12), True

    final_docx3 = os.path.join(OUT_DIR, "Mau_02_TMMH_Thuyet_minh_nhiem_vu.docx")
    final_pdf3 = os.path.join(OUT_DIR, "Mau_02_TMMH_Thuyet_minh_nhiem_vu.pdf")
    doc.save(final_docx3)
    print(f"  [DOCX Saved] {final_docx3}")
    export_docx_to_pdf(final_docx3, final_pdf3)
    return final_docx3

# ==============================================================================
# 4. BIỂU MẪU 4: PHỤ LỤC GIẢI TRÌNH CHI TIẾT CÁC KHOẢN CHI (Khổ Ngang A4 - Filled)
# ==============================================================================
def generate_filled_form4():
    print("\n>>> 4. Đang tạo Phụ lục Khoản chi (Landscape Filled)...")
    doc = create_base_document(orientation="LANDSCAPE", left_mm=18, right_mm=15, top_mm=15, bottom_mm=15)

    add_p(doc, "PHỤ LỤC", bold=True, font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1, line_spacing=1.1)
    add_p(doc, "GIẢI TRÌNH CHI TIẾT CÁC KHOẢN CHI", bold=True, font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1, line_spacing=1.1)
    add_p(doc, "(Kèm theo Thuyết minh nhiệm vụ ứng dụng nhân rộng kết quả nghiên cứu KH&CN)", italic=True, font_size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8, line_spacing=1.1)

    col_widths = [Mm(10), Mm(62), Mm(14), Mm(14), Mm(22), Mm(28), Mm(28), Mm(34), Mm(24), Mm(24)]

    table = doc.add_table(rows=13, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    for c in range(6):
        table.cell(0, c).merge(table.cell(1, c))
    table.cell(0, 9).merge(table.cell(1, 9))
    table.cell(0, 6).merge(table.cell(0, 8))
    table.cell(12, 0).merge(table.cell(12, 1))

    header_row0 = {
        0: ("TT", 11.5),
        1: ("Nội dung", 11.5),
        2: ("ĐVT", 11.5),
        3: ("Số\nlượng", 11),
        4: ("Đơn giá\n(VNĐ)", 11),
        5: ("Thành tiền\n(VNĐ)", 11.5),
        6: ("Trong đó", 11.5),
        9: ("Ghi\nchú", 11.5)
    }
    for c_idx, (txt, sz) in header_row0.items():
        cell = table.cell(0, c_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=40, bottom=40, left=20, right=20)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
        r = p.add_run(txt)
        r.font.name = 'Times New Roman'
        r.font.size, r.font.bold = Pt(sz), True

    header_row1 = {
        6: ("NSSNKH\n(VNĐ)", 10.5),
        7: ("Nguồn tự có của\ncơ quan, tổ chức", 10),
        8: ("Nguồn\nkhác", 10.5)
    }
    for c_idx, (txt, sz) in header_row1.items():
        cell = table.cell(1, c_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=40, bottom=40, left=20, right=20)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
        r = p.add_run(txt)
        r.font.name = 'Times New Roman'
        r.font.size, r.font.bold = Pt(sz), True

    filled_data = [
        # (row_idx, tt, noidung, dvt, sl, dongia, thanhtien, nssnkh, tuco, khac, ghichu, is_bold)
        (2, "I", "Công, thuê khoán chuyên môn", "", "", "", "3.350.000.000", "2.700.000.000", "650.000.000", "-", "Nội dung 1-4", True),
        (3, "1", "Xây dựng đề cương, bộ công cụ & quy trình", "Gói", "1", "350.000.000", "350.000.000", "250.000.000", "100.000.000", "-", "Nội dung 1", False),
        (4, "2", "Điều tra khám sàng lọc tại cộng đồng", "Người", "2.000", "650.000", "1.300.000.000", "1.100.000.000", "200.000.000", "-", "Nội dung 2", False),
        (5, "3", "Khám & chẩn đoán cận lâm sàng chuyên sâu (X-quang, MRI)", "Ca", "700", "2.000.000", "1.400.000.000", "1.100.000.000", "300.000.000", "-", "Nội dung 3", False),
        (6, "4", "Phân tích số liệu và báo cáo dịch tễ học", "Báo cáo", "1", "300.000.000", "300.000.000", "250.000.000", "50.000.000", "-", "Nội dung 4", False),
        (7, "II", "Nguyên vật liệu, năng lượng & Can thiệp lâm sàng", "", "", "", "14.275.000.000", "1.300.000.000", "12.975.000.000", "-", "Nội dung 5", True),
        (8, "1", "Liệu pháp PRP tự thân & Bộ kit li tâm chuẩn", "Ca", "80", "15.000.000", "1.200.000.000", "300.000.000", "900.000.000", "-", "Can thiệp PRP", False),
        (9, "2", "Vật tư phẫu thuật MIS & Thay khớp Robot định vị 3D", "Ca", "50", "261.500.000", "13.075.000.000", "1.000.000.000", "12.075.000.000", "-", "MIS & Robot", False),
        (10, "III", "Thiết bị, máy móc (Khấu hao Robot & Module AI)", "Hệ thống", "1", "-", "-", "-", "-", "-", "Vinmec đối ứng", True),
        (11, "IV", "Chi khác (Khuyến nghị, chuyển giao & Công bố)", "Gói", "1", "800.000.000", "800.000.000", "500.000.000", "300.000.000", "-", "Nội dung 6", True),
    ]

    for r_idx, tt, nd, dvt, sl, dg, ttien, ns, tc, kh, gc, is_b in filled_data:
        vals = [tt, nd, dvt, sl, dg, ttien, ns, tc, kh, gc]
        for c_idx, val in enumerate(vals):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=30, bottom=30, left=20, right=20)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
            if c_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if val:
                r = p.add_run(val)
                r.font.name = 'Times New Roman'
                r.font.size, r.font.bold = Pt(11), is_b

    # Row 12: Tổng cộng
    tot_vals = ["Tổng cộng", "", "", "", "", "18.425.000.000", "4.500.000.000", "13.925.000.000", "0", "100%"]
    for c_idx, val in enumerate(tot_vals):
        try:
            cell = table.cell(12, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=35, bottom=35, left=20, right=20)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if val:
                r = p.add_run(val)
                r.font.name = 'Times New Roman'
                r.font.size, r.font.bold = Pt(11.5), True
        except Exception:
            pass

    for row in table.rows:
        for c_idx, w in enumerate(col_widths):
            try:
                row.cells[c_idx].width = w
            except Exception:
                pass

    final_docx4 = os.path.join(OUT_DIR, "Phu_luc_Giai_trinh_chi_tiet_khoan_chi.docx")
    final_pdf4 = os.path.join(OUT_DIR, "Phu_luc_Giai_trinh_chi_tiet_khoan_chi.pdf")
    doc.save(final_docx4)
    print(f"  [DOCX Saved] {final_docx4}")
    export_docx_to_pdf(final_docx4, final_pdf4)
    return final_docx4

# ==============================================================================
# GOOGLE DRIVE SYNC
# ==============================================================================
def sync_all_to_google_drive(file_map):
    print("\n================================================================================")
    print("BẮT ĐẦU ĐỒNG BỘ 4 BIỂU MẪU LÊN GOOGLE DRIVE THEO FOLDER:")
    print("Folder: https://drive.google.com/drive/folders/1oRPjCHc0FrMnsbRCap7wNnQZZfcTWKyx")
    print("================================================================================")
    creds = service_account.Credentials.from_service_account_file(CREDS_FILE, scopes=['https://www.googleapis.com/auth/drive'])
    drive = build('drive', 'v3', credentials=creds)

    for key, docx_path in file_map.items():
        file_info = DRIVE_FILES[key]
        drive_id = file_info["id"]
        file_name = file_info["name"]
        print(f"\n>>> Đang tải lên file '{file_name}' (ID: {drive_id})...")
        media = MediaFileUpload(
            docx_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            resumable=True
        )
        res = drive.files().update(
            fileId=drive_id,
            media_body=media,
            fields='id, name, modifiedTime, size'
        ).execute()
        print(f"  ✅ ĐÃ ĐỒNG BỘ THÀNH CÔNG: {res.get('name')} (Modified: {res.get('modifiedTime')})")

if __name__ == "__main__":
    f1 = generate_filled_form1()
    f2 = generate_filled_form2()
    f3 = generate_filled_form3()
    f4 = generate_filled_form4()

    sync_all_to_google_drive({
        "form1": f1,
        "form2": f2,
        "form3": f3,
        "form4": f4
    })
    print("\n🎉 HOÀN THÀNH ĐIỀN ĐẦY ĐỦ DỮ LIỆU TỪ PDF VÀO CẢ 4 FORM TRÊN GOOGLE DRIVE!")
