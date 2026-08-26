import os
import sys
import zipfile
import docx
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION_START, WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import win32com.client
import gc
import time

sys.stdout.reconfigure(encoding='utf-8')

BASE_DIR = r"d:\DAOGROUP_WORKSPACE\aVinh"
OUT_DIR = os.path.join(BASE_DIR, "worddata", "chinhthuc")
os.makedirs(OUT_DIR, exist_ok=True)

# Helper: Set Cell Margins (padding) in DXA (1mm = 56.7 dxa)
def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Helper: Set Table Borders (standard single line border)
def set_table_borders(table, color="000000", sz="4"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

# Helper: Add Paragraph
def add_p(doc, text="", bold=False, italic=False, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, 
          space_before=0, space_after=3, line_spacing=1.15, keep_with_next=False, has_tab_dots=False, tab_pos_mm=162, first_line_indent_mm=0, left_indent_mm=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.keep_with_next = keep_with_next
    if first_line_indent_mm > 0:
        p.paragraph_format.first_line_indent = Mm(first_line_indent_mm)
    if left_indent_mm > 0:
        p.paragraph_format.left_indent = Mm(left_indent_mm)
    if has_tab_dots:
        p.paragraph_format.tab_stops.add_tab_stop(Mm(tab_pos_mm), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    if text:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p

# Helper: Create Base Document
def create_base_document(orientation="PORTRAIT", left_mm=30, right_mm=18, top_mm=20, bottom_mm=20):
    doc = docx.Document()
    section = doc.sections[0]
    if orientation == "PORTRAIT":
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT
    else:
        section.page_width = Mm(297)
        section.page_height = Mm(210)
        section.orientation = WD_ORIENT.LANDSCAPE
        
    section.top_margin = Mm(top_mm)
    section.bottom_margin = Mm(bottom_mm)
    section.left_margin = Mm(left_mm)
    section.right_margin = Mm(right_mm)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(14)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(3)
    normal_style.paragraph_format.line_spacing = 1.15
    return doc

# Helper: Inject OpenXML Footnotes
def inject_footnotes_into_docx(in_docx_path, out_docx_path, footnotes_dict):
    fn_items = ""
    for fn_id, fn_text in footnotes_dict.items():
        fn_items += f"""
  <w:footnote w:id="{fn_id}">
    <w:p>
      <w:pPr>
        <w:pStyle w:val="FootnoteText"/>
        <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
          <w:rStyle w:val="FootnoteReference"/>
          <w:vertAlign w:val="superscript"/>
          <w:sz w:val="20"/>
        </w:rPr>
        <w:footnoteRef/>
      </w:r>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
          <w:i/>
          <w:sz w:val="22"/>
        </w:rPr>
        <w:t xml:space="preserve"> {fn_text}</w:t>
      </w:r>
    </w:p>
  </w:footnote>"""

    footnotes_xml_content = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:footnote w:type="separator" w:id="-1">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>{fn_items}
</w:footnotes>""".encode('utf-8')

    staging_path = out_docx_path + ".tmp"
    with zipfile.ZipFile(in_docx_path, 'r') as zin:
        with zipfile.ZipFile(staging_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    ct_str = data.decode('utf-8')
                    if 'footnotes.xml' not in ct_str:
                        fn_override = '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                        ct_str = ct_str.replace('</Types>', f'{fn_override}</Types>')
                        data = ct_str.encode('utf-8')
                elif item.filename == 'word/_rels/document.xml.rels':
                    rels_str = data.decode('utf-8')
                    if 'footnotes.xml' not in rels_str:
                        fn_rel = '<Relationship Id="rIdFootnotes" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
                        rels_str = rels_str.replace('</Relationships>', f'{fn_rel}</Relationships>')
                        data = rels_str.encode('utf-8')
                zout.writestr(item, data)
            zout.writestr('word/footnotes.xml', footnotes_xml_content)
    for _ in range(5):
        try:
            if os.path.exists(out_docx_path):
                try:
                    os.remove(out_docx_path)
                except Exception:
                    pass
            os.replace(staging_path, out_docx_path)
            break
        except Exception:
            time.sleep(0.5)

# Helper: Export DOCX to PDF using Word COM
def export_docx_to_pdf(docx_path, pdf_path):
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc_word = None
    try:
        doc_word = word.Documents.Open(os.path.abspath(docx_path))
        doc_word.ExportAsFixedFormat(os.path.abspath(pdf_path), 17)
        doc_word.Close(False)
        doc_word = None
        print(f"  [PDF Exported] {pdf_path}")
    finally:
        word.Quit()
        if doc_word is not None:
            try:
                doc_word.Close(False)
            except Exception:
                pass
            del doc_word
        del word
        gc.collect()
        time.sleep(0.3)

# Helper: Add Footnote Reference XML
def add_footnote_reference(paragraph, footnote_id=1, font_size_half_pt=22):
    fn_ref = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        f'<w:rStyle w:val="FootnoteReference"/>'
        f'<w:vertAlign w:val="superscript"/>'
        f'<w:b/>'
        f'<w:sz w:val="{font_size_half_pt}"/>'
        f'</w:rPr>'
        f'<w:footnoteReference w:id="{footnote_id}"/>'
        f'</w:r>'
    )
    paragraph._p.append(fn_ref)

print("================================================================================")
print("BẮT ĐẦU SINH CÁC BIỂU MẪU CHUẨN TỪ TRANG 8 ĐẾN TRANG 15 VÀO THƯ MỤC CHÍNH THỨC")
print("================================================================================")

# ==============================================================================
# 1. BIỂU MẪU 1: PHIẾU ĐẶT HÀNG NHIỆM VỤ KH&CN (Mẫu I.01-ĐHNV - Trang 8-9)
# ==============================================================================
print("\n>>> 1. Đang tạo Mẫu I.01-ĐHNV: Mau_I01_DHNV_Phieu_dat_hang...")
doc1 = create_base_document(orientation="PORTRAIT", left_mm=30, right_mm=18, top_mm=20, bottom_mm=20)

add_p(doc1, "PHỤ LỤC II. BIỂU MẪU", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1, line_spacing=1.15)
add_p(doc1, "Mẫu I.01-ĐHNV", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=4, line_spacing=1.15)

top_table1 = doc1.add_table(rows=1, cols=2)
top_table1.alignment = WD_TABLE_ALIGNMENT.CENTER
top_table1.autofit = False

c_left = top_table1.cell(0, 0)
c_right = top_table1.cell(0, 1)
c_left.width = Mm(52)
c_right.width = Mm(110)

set_cell_margins(c_left, top=0, bottom=0, left=0, right=0)
p_org1 = c_left.paragraphs[0]
p_org1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_org1.paragraph_format.space_before, p_org1.paragraph_format.space_after, p_org1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
r = p_org1.add_run("TÊN CƠ QUAN CHỦ QUẢN")
r.font.size = Pt(12)

p_org2 = c_left.add_paragraph()
p_org2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_org2.paragraph_format.space_before, p_org2.paragraph_format.space_after, p_org2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = p_org2.add_run("TÊN TỔ CHỨC")
r.font.size, r.font.bold = Pt(12), True

p_org3 = c_left.add_paragraph()
p_org3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_org3.paragraph_format.space_before, p_org3.paragraph_format.space_after, p_org3.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
r = p_org3.add_run("______")
r.font.size, r.font.bold = Pt(11), True

set_cell_margins(c_right, top=0, bottom=0, left=0, right=0)
p_m1 = c_right.paragraphs[0]
p_m1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_m1.paragraph_format.space_before, p_m1.paragraph_format.space_after, p_m1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
r = p_m1.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
r.font.size, r.font.bold = Pt(12), True

p_m2 = c_right.add_paragraph()
p_m2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_m2.paragraph_format.space_before, p_m2.paragraph_format.space_after, p_m2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = p_m2.add_run("Độc lập - Tự do - Hạnh phúc")
r.font.size, r.font.bold = Pt(13), True

p_m3 = c_right.add_paragraph()
p_m3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_m3.paragraph_format.space_before, p_m3.paragraph_format.space_after, p_m3.paragraph_format.line_spacing = Pt(0), Pt(3), 1.0
r = p_m3.add_run("_________________")
r.font.size, r.font.bold = Pt(11), True

p_d = c_right.add_paragraph()
p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_d.paragraph_format.space_before, p_d.paragraph_format.space_after, p_d.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r_d1 = p_d.add_run("........")
r_d1.font.size = Pt(13)
r_d2 = p_d.add_run(", ngày… tháng…  năm 20....")
r_d2.font.size, r_d2.font.italic = Pt(13), True

p_title1 = add_p(doc1, "", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=14, space_after=14, line_spacing=1.15)
r_t = p_title1.add_run("PHIẾU ĐẶT HÀNG NHIỆM VỤ KHOA HỌC VÀ CÔNG NGHỆ")
r_t.font.size, r_t.font.bold = Pt(14), True
add_footnote_reference(p_title1, footnote_id=1, font_size_half_pt=24)

# Page 8 body items
add_p(doc1, "1. Tên nhiệm vụ:\t", space_after=2, has_tab_dots=True)
add_p(doc1, "\t", space_after=4, has_tab_dots=True)
add_p(doc1, "2. Thuộc lĩnh vực:", space_after=2)

tbl_field = doc1.add_table(rows=4, cols=2)
tbl_field.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_field.autofit = False

fields_data = [
    ("- Khoa học tự nhiên", "- Khoa học kỹ thuật và công nghệ"),
    ("- Khoa học y, dược", "- Khoa học nông nghiệp"),
    ("- Khoa học xã hội", "- Khoa học nhân văn"),
    ("- Công nghệ chiến lược", "")
]
for row_idx, (c1, c2) in enumerate(fields_data):
    cell1 = tbl_field.cell(row_idx, 0)
    cell2 = tbl_field.cell(row_idx, 1)
    cell1.width, cell2.width = Mm(81), Mm(81)
    for cell, txt in [(cell1, c1), (cell2, c2)]:
        set_cell_margins(cell, top=0, bottom=0, left=50, right=50)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(2), 1.2
        p.paragraph_format.tab_stops.add_tab_stop(Mm(76), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        if txt:
            r1 = p.add_run(txt + "\t")
            r1.font.size = Pt(14)
            r2 = p.add_run("☐")
            r2.font.name, r2.font.size = 'Segoe UI Symbol', Pt(14)

add_p(doc1, "3. Thuộc loại hình nhiệm vụ", space_before=4, space_after=2)
tbl_type = doc1.add_table(rows=2, cols=2)
tbl_type.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_type.autofit = False

types_data = [
    ("- Nghiên cứu cơ bản", "- Nghiên cứu ứng dụng"),
    ("- Phát triển công nghệ", "- Phát triển giải pháp xã hội")
]
for row_idx, (c1, c2) in enumerate(types_data):
    cell1 = tbl_type.cell(row_idx, 0)
    cell2 = tbl_type.cell(row_idx, 1)
    cell1.width, cell2.width = Mm(81), Mm(81)
    for cell, txt in [(cell1, c1), (cell2, c2)]:
        set_cell_margins(cell, top=0, bottom=0, left=50, right=50)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(2), 1.2
        p.paragraph_format.tab_stops.add_tab_stop(Mm(76), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        r1 = p.add_run(txt + "\t")
        r1.font.size = Pt(14)
        r2 = p.add_run("☐")
        r2.font.name, r2.font.size = 'Segoe UI Symbol', Pt(14)

add_p(doc1, "4. Căn cứ đặt hàng nhiệm vụ:\t", space_before=4, space_after=2, has_tab_dots=True)
add_p(doc1, "\t", space_after=4, has_tab_dots=True)
add_p(doc1, "5. Tính cấp thiết của nhiệm vụ:\t", space_after=2, has_tab_dots=True)
add_p(doc1, "\t", space_after=4, has_tab_dots=True)
add_p(doc1, "6. Mục tiêu:\t", space_after=2, has_tab_dots=True)
add_p(doc1, "\t", space_after=4, has_tab_dots=True)
add_p(doc1, "7. Dự kiến các nội dung chính cần thực hiện:\t", space_after=2, has_tab_dots=True)
add_p(doc1, "\t", space_after=4, has_tab_dots=True)
add_p(doc1, "8. Dự kiến các kết quả thực hiện nhiệm vụ và các chỉ tiêu cần đạt:\t", space_after=2, has_tab_dots=True)
add_p(doc1, "\t", space_after=4, has_tab_dots=True)
add_p(doc1, "9. Hiệu quả và tác động của kết quả thực hiện nhiệm vụ:\t", space_after=2, has_tab_dots=True)
add_p(doc1, "\t", space_after=4, has_tab_dots=True)

# Page 9 body items: ĐÚNG THỨ TỰ BẢN GỐC (10 -> 11 -> - Kinh phí NSNN -> - Kinh phí nguồn khác)
add_p(doc1, "10.  Dự kiến kinh phí thực hiện:\t", space_after=2, keep_with_next=True, has_tab_dots=True)
add_p(doc1, "11.  Dự kiến thời gian thực hiện:\t", space_after=2, keep_with_next=True, has_tab_dots=True)
add_p(doc1, "- Kinh phí hỗ trợ từ NSNN:\t", space_after=2, keep_with_next=True, has_tab_dots=True)
add_p(doc1, "- Kinh phí từ nguồn khác:\t", space_after=3, has_tab_dots=True)

add_p(doc1, "12. Đề xuất tổ chức được xét giao trực tiếp (nếu có):\t", space_after=2, has_tab_dots=True)
add_p(doc1, "\t", space_after=4, has_tab_dots=True)

add_p(doc1, "13. Tiếp nhận và phương án tổ chức quản lý, sử dụng kết quả của nhiệm vụ của", space_after=1)
add_p(doc1, "cơ quan đặt hàng: ", space_after=2)
add_p(doc1, "(Nêu rõ việc Nhà nước sẽ nắm giữ quyền quản lý, sử dụng, quyền sở hữu kết quả trong trường hợp nhà nước có yêu cầu tiếp nhận kết quả)", italic=True, space_after=2)
add_p(doc1, "\t", space_after=2, has_tab_dots=True)
add_p(doc1, "\t", space_after=4, has_tab_dots=True)

add_p(doc1, "14. Danh mục tài liệu tham khảo:", space_after=4)

add_p(doc1, "15. Thông tin liên hệ:", space_after=2)
add_p(doc1, "Tên tổ chức:\t", space_after=2, has_tab_dots=True)
add_p(doc1, "Đại diện tổ chức:\t", space_after=2, has_tab_dots=True)
add_p(doc1, "Mã định danh điện tử của tổ chức:\t", space_after=2, has_tab_dots=True)
add_p(doc1, "Điện thoại:\t", space_after=2, has_tab_dots=True)
add_p(doc1, "Email:\t", space_after=2, has_tab_dots=True)
add_p(doc1, "Địa chỉ liên hệ:\t", space_after=14, has_tab_dots=True)

tbl_sig1 = doc1.add_table(rows=1, cols=2)
tbl_sig1.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_sig1.autofit = False
c_sig_l1 = tbl_sig1.cell(0, 0)
c_sig_r1 = tbl_sig1.cell(0, 1)
c_sig_l1.width, c_sig_r1.width = Mm(70), Mm(92)

set_cell_margins(c_sig_r1, top=0, bottom=0, left=0, right=0)
p_s1 = c_sig_r1.paragraphs[0]
p_s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_s1.paragraph_format.space_before, p_s1.paragraph_format.space_after, p_s1.paragraph_format.line_spacing = Pt(8), Pt(2), 1.15
r = p_s1.add_run("ĐẠI DIỆN TỔ CHỨC ĐẶT HÀNG")
r.font.size, r.font.bold = Pt(14), True

p_s2 = c_sig_r1.add_paragraph()
p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_s2.paragraph_format.space_before, p_s2.paragraph_format.space_after, p_s2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = p_s2.add_run("(Họ, tên và chữ ký - đóng dấu)")
r.font.size, r.font.italic = Pt(14), True

temp_docx1 = os.path.join(OUT_DIR, "temp_form1.docx")
final_docx1 = os.path.join(OUT_DIR, "Mau_I01_DHNV_Phieu_dat_hang.docx")
final_pdf1 = os.path.join(OUT_DIR, "Mau_I01_DHNV_Phieu_dat_hang.pdf")

doc1.save(temp_docx1)
inject_footnotes_into_docx(temp_docx1, final_docx1, {1: "Phiếu đặt hàng được trình bày không quá 10 trang giấy khổ A4."})
os.remove(temp_docx1)
print(f"  [DOCX Saved] {final_docx1}")
export_docx_to_pdf(final_docx1, final_pdf1)


# ==============================================================================
# 2. BIỂU MẪU 2: TỜ TRÌNH PHÊ DUYỆT (Mẫu 02 -TTMH - Trang 10-11)
# ==============================================================================
print("\n>>> 2. Đang tạo Mẫu 02 -TTMH: Mau_02_TTMH_To_trinh_phe_duyet...")
doc2 = create_base_document(orientation="PORTRAIT", left_mm=30, right_mm=18, top_mm=20, bottom_mm=20)

add_p(doc2, "Mẫu 02 -TTMH", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=4, line_spacing=1.15)

top_table2 = doc2.add_table(rows=1, cols=2)
top_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
top_table2.autofit = False

c_left2 = top_table2.cell(0, 0)
c_right2 = top_table2.cell(0, 1)
c_left2.width, c_right2.width = Mm(65), Mm(97)

set_cell_margins(c_left2, top=0, bottom=0, left=0, right=0)
p_t_l1 = c_left2.paragraphs[0]
p_t_l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t_l1.paragraph_format.space_before, p_t_l1.paragraph_format.space_after, p_t_l1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
r = p_t_l1.add_run("CƠ QUAN TRÌNH")
r.font.size = Pt(12)

p_t_l2 = c_left2.add_paragraph()
p_t_l2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t_l2.paragraph_format.space_before, p_t_l2.paragraph_format.space_after, p_t_l2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = p_t_l2.add_run("PHÊ DUYỆT")
r.font.size, r.font.bold = Pt(12), True

p_t_l3 = c_left2.add_paragraph()
p_t_l3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t_l3.paragraph_format.space_before, p_t_l3.paragraph_format.space_after, p_t_l3.paragraph_format.line_spacing = Pt(0), Pt(2), 1.0
r = p_t_l3.add_run("-------")
r.font.size, r.font.bold = Pt(11), True

p_num2 = c_left2.add_paragraph()
p_num2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_num2.paragraph_format.space_before, p_num2.paragraph_format.space_after, p_num2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = p_num2.add_run("Số:……./TTr-…...")
r.font.size = Pt(13)

set_cell_margins(c_right2, top=0, bottom=0, left=0, right=0)
p_t_r1 = c_right2.paragraphs[0]
p_t_r1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t_r1.paragraph_format.space_before, p_t_r1.paragraph_format.space_after, p_t_r1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
r = p_t_r1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
r.font.size, r.font.bold = Pt(12), True

p_t_r2 = c_right2.add_paragraph()
p_t_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t_r2.paragraph_format.space_before, p_t_r2.paragraph_format.space_after, p_t_r2.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
r = p_t_r2.add_run("Độc lập - Tự do - Hạnh phúc")
r.font.size, r.font.bold = Pt(13), True

p_date2 = c_right2.add_paragraph()
p_date2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date2.paragraph_format.space_before, p_date2.paragraph_format.space_after, p_date2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = p_date2.add_run("……., ngày... tháng... năm…..")
r.font.size, r.font.italic = Pt(13), True

add_p(doc2, "TỜ TRÌNH", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=2, line_spacing=1.15)
add_p(doc2, "Về việc phê duyệt, hỗ trợ nhiệm vụ ứng dụng, chuyển giao tiến bộ\nkhoa học và công nghệ vào đời sống và sản xuất", 
      bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1, line_spacing=1.2)

p_subj_line2 = add_p(doc2, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=10, line_spacing=1.0)
r_line2 = p_subj_line2.add_run("____________________")
r_line2.font.name = 'Times New Roman'
r_line2.font.size = Pt(11)
r_line2.font.bold = True

add_p(doc2, "Kính gửi:   ", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=2, left_indent_mm=34)
add_p(doc2, "- Sở Khoa học và Công nghệ;", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=1, left_indent_mm=56)
add_p(doc2, "- Sở Tài chính.", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, left_indent_mm=56)

add_p(doc2, "Căn cứ Luật Ngân sách nhà nước ngày 25 tháng 6 năm 2025;", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)
add_p(doc2, "Căn cứ Luật Quản lý, sử dụng tài sản công ngày 21 tháng 6 năm 2017;", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)
add_p(doc2, "Căn cứ Nghị định số 104/2026/NĐ-CP ngày 31 tháng 3 năm 2026 của Chính phủ quy định việc lập dự toán, quản lý, sử dụng và quyết toán chi thường xuyên để thực hiện các nhiệm vụ quy định tại Điều 40 Luật Ngân sách nhà nước;", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)
add_p(doc2, "Căn cứ các căn cứ pháp lý, văn bản liên quan khác (nếu có);", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=12.7)

p_intro2 = add_p(doc2, "", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=12.7)
r1 = p_intro2.add_run("(Cơ quan trình phê duyệt)")
r1.font.name, r1.font.size, r1.font.italic = 'Times New Roman', Pt(14), True
r2 = p_intro2.add_run("…. Sở Khoa học và Công nghệ, Sở Tài chính xem xét thẩm định, phê duyệt hỗ trợ nhiệm vụ ứng dụng, chuyển giao tiến bộ khoa học và công nghệ vào đời sống và sản xuất năm 2027 cho ")
r2.font.name, r2.font.size = 'Times New Roman', Pt(14)
r3 = p_intro2.add_run("(Cơ quan trình phê duyệt)")
r3.font.name, r3.font.size, r3.font.italic = 'Times New Roman', Pt(14), True
r4 = p_intro2.add_run("….với các nội dung chủ yếu như sau:")
r4.font.name, r4.font.size = 'Times New Roman', Pt(14)

add_p(doc2, "I. THÔNG TIN CHUNG", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=2)
add_p(doc2, "1. Tên cơ quan, đơn vị đề xuất: ", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
add_p(doc2, "2. Tên nhiệm vụ:", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
add_p(doc2, "3. Dự kiến kinh phí:…. triệu đồng. Trong đó:", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
add_p(doc2, "- Nguồn ngân sách nhà nước: .... triệu đồng", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)
add_p(doc2, "- Nguồn khác (nếu có): .... triệu đồng", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)

p_time = add_p(doc2, "4. Thời gian thực hiện nhiệm vụ", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
add_footnote_reference(p_time, footnote_id=1, font_size_half_pt=22)
r_dot = p_time.add_run(".")
r_dot.font.size = Pt(14)

add_p(doc2, "5. Các nội dung khác (nếu có).", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

add_p(doc2, "II. NGUỒN KINH PHÍ THỰC HIỆN", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=2)
add_p(doc2, "Đề nghị hỗ trợ từ nguồn ngân sách sự nghiệp KHCN, ĐMST&CĐS năm 2027 của tỉnh Hưng Yên", 
      font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=12.7)

add_p(doc2, "III. HỒ SƠ TÀI LIỆU KÈM THEO", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=2)
add_p(doc2, "- Thuyết minh chi tiết nhiệm vụ kèm theo dự toán;", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)
add_p(doc2, "- Báo giá và các tài liệu liên quan khác kèm theo.", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=12.7)

p_outro2 = add_p(doc2, "", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=4, space_after=12, first_line_indent_mm=12.7)
r_o1 = p_outro2.add_run("(Cơ quan trình phê duyệt)")
r_o1.font.name, r_o1.font.size, r_o1.font.italic = 'Times New Roman', Pt(14), True
r_o2 = p_outro2.add_run("…. kính trình Sở Khoa học và Công nghệ, Sở Tài chính xem xét thẩm định, trình UBND tỉnh phê duyệt./.")
r_o2.font.name, r_o2.font.size = 'Times New Roman', Pt(14)

tbl_sig2 = doc2.add_table(rows=1, cols=2)
tbl_sig2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_sig2.autofit = False

c_noi_nhan = tbl_sig2.cell(0, 0)
c_ky_ten = tbl_sig2.cell(0, 1)
c_noi_nhan.width, c_ky_ten.width = Mm(75), Mm(87)

set_cell_margins(c_noi_nhan, top=0, bottom=0, left=0, right=0)
p_nn1 = c_noi_nhan.paragraphs[0]
p_nn1.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_nn1.paragraph_format.space_before, p_nn1.paragraph_format.space_after, p_nn1.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
r = p_nn1.add_run("Nơi nhận:\n")
r.font.size, r.font.bold, r.font.italic = Pt(12), True, True
for line in ["- Như trên;", "- Các cơ quan có liên quan;", "- Lưu:..."]:
    r = p_nn1.add_run(line + "\n")
    r.font.size = Pt(11)

set_cell_margins(c_ky_ten, top=0, bottom=0, left=0, right=0)
p_k1 = c_ky_ten.paragraphs[0]
p_k1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_k1.paragraph_format.space_before, p_k1.paragraph_format.space_after, p_k1.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
r = p_k1.add_run("CƠ QUAN TRÌNH PHÊ DUYỆT")
r.font.size, r.font.bold = Pt(14), True

p_k2 = c_ky_ten.add_paragraph()
p_k2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_k2.paragraph_format.space_before, p_k2.paragraph_format.space_after, p_k2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = p_k2.add_run("(Ký, ghi rõ họ tên, chức vụ và\nđóng dấu)")
r.font.size, r.font.italic = Pt(13), True

temp_docx2 = os.path.join(OUT_DIR, "temp_form2.docx")
final_docx2 = os.path.join(OUT_DIR, "Mau_02_TTMH_To_trinh_phe_duyet.docx")
final_pdf2 = os.path.join(OUT_DIR, "Mau_02_TTMH_To_trinh_phe_duyet.pdf")

doc2.save(temp_docx2)
inject_footnotes_into_docx(temp_docx2, final_docx2, {1: "Trong trường hợp nhiệm vụ có thời gian thực hiện trên 01 năm thì cần xác định dự toán kinh phí thực hiện trong từng năm."})
os.remove(temp_docx2)
print(f"  [DOCX Saved] {final_docx2}")
export_docx_to_pdf(final_docx2, final_pdf2)


# ==============================================================================
# 3. BIỂU MẪU 3: THUYẾT MINH NHIỆM VỤ (Mẫu 02 -TMMH - Trang 12-14)
# ==============================================================================
print("\n>>> 3. Đang tạo Mẫu 02 -TMMH: Mau_02_TMMH_Thuyet_minh_nhiem_vu...")
doc3 = create_base_document(orientation="PORTRAIT", left_mm=30, right_mm=18, top_mm=20, bottom_mm=20)

# --- TRANG BÌA (Page 12) ---
add_p(doc3, "Mẫu 02 -TMMH", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=14, line_spacing=1.15)
add_p(doc3, "………..(ĐƠN VỊ QUẢN LÝ)", bold=True, font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2, line_spacing=1.15)
add_p(doc3, "………….. (ĐƠN VỊ THỰC HIỆN)", bold=True, font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=40, line_spacing=1.15)

add_p(doc3, "THUYẾT MINH NHIỆM VỤ", bold=True, font_size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4, line_spacing=1.2)
add_p(doc3, "ỨNG DỤNG, NHÂN RỘNG KẾT QUẢ NGHIÊN CỨU\nKHOA HỌC VÀ CÔNG NGHỆ", bold=True, font_size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=40, line_spacing=1.25)

add_p(doc3, "Tên nhiệm vụ:……………………………………………………………….", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, line_spacing=1.2)
add_p(doc3, "Ứng dụng, nhân rộng kết quả của nhiệm vụ KH&CN cấp tỉnh:", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=3, line_spacing=1.2)
add_p(doc3, "“………..”", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6, line_spacing=1.2)
add_p(doc3, "Đơn vị quản lý:…………", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, line_spacing=1.2)
add_p(doc3, "Đơn vị thực hiện:………………", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, line_spacing=1.2)
add_p(doc3, "Thời gian thực hiện: Năm 2027", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=50, line_spacing=1.2)
add_p(doc3, "Hưng Yên, năm….", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0, line_spacing=1.2)

# Page Break sang Thân thuyết minh
doc3.add_page_break()

# --- THÂN THUYẾT MINH (Pages 13 - 14) ---
add_p(doc3, "THUYẾT MINH NHIỆM VỤ", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2, line_spacing=1.15)
add_p(doc3, "ỨNG DỤNG, NHÂN RỘNG KẾT QUẢ CỦA NHIỆM VỤ KH&CN", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=10, line_spacing=1.15)

add_p(doc3, "I. THÔNG TIN CHUNG", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=3)
add_p(doc3, "1. Tên nhiệm vụ:", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc3, "- Ứng dụng, nhân rộng kết quả của nhiệm vụ KH&CN cấp tỉnh năm ………: “………………………………………………………………………..…..”", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
add_p(doc3, "Mã số (nếu có): ……………………………..", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

add_p(doc3, "2. Tính cấp thiết cần triển khai thực hiện nhiệm vụ", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc3, "Lý giải vì sao phải thực hiện nhiệm vụ này, sự cần thiết phải triển khai. Phân tích thực trạng sản xuất trong lĩnh vực đó và những hiệu quả trong thời gian vừa qua. Nhiệm vụ được triển khai sẽ góp phần nâng cao hiệu quả sản xuất, thay đổi tập quán, phương thức canh tác, chuẩn hóa các quy trình sản xuất….", 
      italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=12.7)

add_p(doc3, "3. Quy mô, địa điểm, thời gian thực hiện", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc3, "- Quy mô: Diện tích, số lượng…", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc3, "- Địa điểm:", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc3, "- Thời gian: … tháng, từ tháng… đến tháng…năm …", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_line_indent_mm=12.7)

add_p(doc3, "4. Kinh phí thực hiện", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc3, "Tổng kinh phí:…. Triệu đồng", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc3, "Trong đó:", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc3, "- Ngân sách SNKH:…. Triệu đồng", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc3, "- Nguồn khác: …. Triệu đồng", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_line_indent_mm=20)

add_p(doc3, "5. Đơn vị phối hợp thực hiện, đơn vị hướng dẫn kỹ thuật", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
add_p(doc3, "6. Đơn vị quản lý", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

add_p(doc3, "II. MỤC TIÊU, NỘI DUNG, PHƯƠNG ÁN TỔ CHỨC VÀ TIẾN ĐỘ THỰC HIỆN", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=3)
add_p(doc3, "1. Mục tiêu", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc3, "1.1. Mục tiêu chung", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc3, "1.2. Mục tiêu cụ thể", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc3, "- Mục tiêu cụ thể 1", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc3, "- Mục tiêu cụ thể 2", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc3, "- Mục tiêu cụ thể 3", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_line_indent_mm=20)

add_p(doc3, "2. Nội dung thực hiện (Luận giải những công việc cần thực hiện để hoàn thành các mục tiêu cụ thể, mục tiêu chung đã đề ra ở trên, đảm bảo tính thực tiễn, phù hợp với mục tiêu)", 
      italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3)
add_p(doc3, "2.1. Nội dung 1", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc3, "- Công việc 1", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc3, "- Công việc 2", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc3, "- Công việc 3", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3, first_line_indent_mm=20)

add_p(doc3, "2.2. Nội dung 2", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc3, "- Công việc 1", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc3, "- Công việc 2", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc3, "- Công việc 3", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3, first_line_indent_mm=20)

add_p(doc3, "2.3. Nội dung 3", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc3, "- Công việc 1", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc3, "- Công việc 2", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc3, "- Công việc 3", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc3, "……………", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_line_indent_mm=12.7)

add_p(doc3, "3. Phương án tổ chức và tiến độ thực hiện", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc3, "3.1. Phương án tổ chức", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc3, "3.2. Tiến độ thực hiện", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_line_indent_mm=12.7)

headers_sched = ["Stt", "Nội dung công việc", "Kết quả đạt được", "Tiến độ thời gian", "Đơn vị, cá nhân thực hiện", "Ghi chú"]
col_widths_sched = [Mm(14), Mm(44), Mm(36), Mm(26), Mm(30), Mm(12)]

tbl_schedule = doc3.add_table(rows=5, cols=6)
tbl_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_schedule.autofit = False
set_table_borders(tbl_schedule)

for c_idx, (h_text, w) in enumerate(zip(headers_sched, col_widths_sched)):
    cell = tbl_schedule.cell(0, c_idx)
    cell.width = w
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=80, bottom=80, left=60, right=60)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.1
    r = p.add_run(h_text)
    r.font.size, r.font.bold = Pt(13), True

for r_idx in range(1, 5):
    for c_idx, w in enumerate(col_widths_sched):
        cell = tbl_schedule.cell(r_idx, c_idx)
        cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=100, bottom=100, left=60, right=60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 3, 5] else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.1
        if c_idx == 0:
            r = p.add_run(str(r_idx))
            r.font.size = Pt(13)

add_p(doc3, "III. DỰ KIẾN SẢN PHẨM, PHÂN TÍCH HIỆU QUẢ, KẾ HOẠCH NHÂN RỘNG", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=3)
add_p(doc3, "1. Dự kiến sản phẩm của nhiệm vụ", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc3, "2. Phân tích hiệu quả (hiệu quả kinh tế, xã hội, môi trường…)", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc3, "3. Kế hoạch triển khai nhân rộng nhiệm vụ", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

add_p(doc3, "Hưng Yên, ngày …. tháng…. năm …..", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=4, space_after=8)

tbl_sig3 = doc3.add_table(rows=1, cols=2)
tbl_sig3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_sig3.autofit = False

c_thuc_hien = tbl_sig3.cell(0, 0)
c_quan_ly = tbl_sig3.cell(0, 1)
c_thuc_hien.width, c_quan_ly.width = Mm(81), Mm(81)

set_cell_margins(c_thuc_hien, top=0, bottom=0, left=0, right=0)
p_th1 = c_thuc_hien.paragraphs[0]
p_th1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_th1.paragraph_format.space_before, p_th1.paragraph_format.space_after, p_th1.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
r = p_th1.add_run("ĐƠN VỊ THỰC HIỆN")
r.font.size, r.font.bold = Pt(14), True

p_th2 = c_thuc_hien.add_paragraph()
p_th2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_th2.paragraph_format.space_before, p_th2.paragraph_format.space_after, p_th2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = p_th2.add_run("(Ký và đóng dấu)")
r.font.size, r.font.italic = Pt(14), True

set_cell_margins(c_quan_ly, top=0, bottom=0, left=0, right=0)
p_ql1 = c_quan_ly.paragraphs[0]
p_ql1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ql1.paragraph_format.space_before, p_ql1.paragraph_format.space_after, p_ql1.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
r = p_ql1.add_run("ĐƠN VỊ QUẢN LÝ")
r.font.size, r.font.bold = Pt(14), True

p_ql2 = c_quan_ly.add_paragraph()
p_ql2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ql2.paragraph_format.space_before, p_ql2.paragraph_format.space_after, p_ql2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = p_ql2.add_run("(Ký và đóng dấu)")
r.font.size, r.font.italic = Pt(14), True

final_docx3 = os.path.join(OUT_DIR, "Mau_02_TMMH_Thuyet_minh_nhiem_vu.docx")
final_pdf3 = os.path.join(OUT_DIR, "Mau_02_TMMH_Thuyet_minh_nhiem_vu.pdf")
doc3.save(final_docx3)
print(f"  [DOCX Saved] {final_docx3}")
export_docx_to_pdf(final_docx3, final_pdf3)


# ==============================================================================
# ==============================================================================
# 4. BIỂU MẪU 4: PHỤ LỤC GIẢI TRÌNH CHI TIẾT CÁC KHOẢN CHI (Trang 15 - LANDSCAPE 1 TRANG DUY NHẤT)
# ==============================================================================
print("\n>>> 4. Đang tạo Phụ lục: Phu_luc_Giai_trinh_chi_tiet_khoan_chi (Landscape 1 Trang)...")

def build_landscape_phuluc(doc):
    # Title
    add_p(doc, "PHỤ LỤC", bold=True, font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1, line_spacing=1.1)
    add_p(doc, "GIẢI TRÌNH CHI TIẾT CÁC KHOẢN CHI", bold=True, font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1, line_spacing=1.1)
    add_p(doc, "(Kèm theo Thuyết minh nhiệm vụ ứng dụng nhân rộng kết quả nghiên cứu KH&CN)", italic=True, font_size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8, line_spacing=1.1)

    col_widths = [Mm(12), Mm(64), Mm(16), Mm(18), Mm(22), Mm(26), Mm(26), Mm(34), Mm(24), Mm(18)]

    table = doc.add_table(rows=13, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    # 1. Vertical merges for Header
    for c in range(6):
        table.cell(0, c).merge(table.cell(1, c))
    table.cell(0, 9).merge(table.cell(1, 9))

    # 2. Horizontal merge for 'Trong đó' (cols 6, 7, 8 in row 0)
    table.cell(0, 6).merge(table.cell(0, 8))

    # 3. Merge for 'Tổng cộng' in last row
    table.cell(12, 0).merge(table.cell(12, 1))

    # Set Header Content Row 0
    header_row0 = {
        0: ("TT", 12),
        1: ("Nội dung", 12),
        2: ("ĐVT", 12),
        3: ("Số\nlượng", 11.5),
        4: ("Đơn giá", 12),
        5: ("Thành tiền", 12),
        6: ("Trong đó", 12),
        9: ("Ghi\nchú", 12)
    }
    for c_idx, (txt, sz) in header_row0.items():
        cell = table.cell(0, c_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=40, bottom=40, left=30, right=30)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
        r = p.add_run(txt)
        r.font.name = 'Times New Roman'
        r.font.size, r.font.bold = Pt(sz), True

    # Set Header Content Row 1 (sub-headers under 'Trong đó')
    header_row1 = {
        6: ("NSSNKH", 11),
        7: ("Nguồn tự có\ncủa cơ quan,\ntổ chức", 10.5),
        8: ("Nguồn\nkhác", 11)
    }
    for c_idx, (txt, sz) in header_row1.items():
        cell = table.cell(1, c_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=40, bottom=40, left=30, right=30)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
        r = p.add_run(txt)
        r.font.name = 'Times New Roman'
        r.font.size, r.font.bold = Pt(sz), True

    # Data Rows
    row_data = [
        (2, "I", "Công, thuê khoán chuyên môn", True),
        (3, "1", "", False),
        (4, "2", "", False),
        (5, "…", "", False),
        (6, "II", "Nguyên vật liệu, năng lượng", True),
        (7, "1", "", False),
        (8, "2", "", False),
        (9, "…", "", False),
        (10, "III", "Thiết bị, máy móc", True),
        (11, "IV", "Chi khác", True),
    ]

    for r_idx, tt_val, content_val, is_bold in row_data:
        for c_idx in range(10):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=35, bottom=35, left=30, right=30)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if tt_val:
                    r = p.add_run(tt_val)
                    r.font.name = 'Times New Roman'
                    r.font.size, r.font.bold = Pt(12), is_bold
            elif c_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if content_val:
                    r = p.add_run(content_val)
                    r.font.name = 'Times New Roman'
                    r.font.size, r.font.bold = Pt(12), is_bold
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 12: Tổng cộng
    for c_idx in range(10):
        try:
            cell = table.cell(12, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=35, bottom=35, left=30, right=30)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
            if c_idx in (0, 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not p.text:
                    r = p.add_run("Tổng cộng")
                    r.font.name = 'Times New Roman'
                    r.font.size, r.font.bold = Pt(12), True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    # Set column widths across all rows and cells
    for row in table.rows:
        for c_idx, w in enumerate(col_widths):
            try:
                row.cells[c_idx].width = w
            except Exception:
                pass

    return table

doc4 = create_base_document(orientation="LANDSCAPE", left_mm=18, right_mm=15, top_mm=15, bottom_mm=15)
build_landscape_phuluc(doc4)

final_docx4 = os.path.join(OUT_DIR, "Phu_luc_Giai_trinh_chi_tiet_khoan_chi.docx")
final_pdf4 = os.path.join(OUT_DIR, "Phu_luc_Giai_trinh_chi_tiet_khoan_chi.pdf")
doc4.save(final_docx4)
print(f"  [DOCX Saved] {final_docx4}")
export_docx_to_pdf(final_docx4, final_pdf4)



# ==============================================================================
# 5. TÀI LIỆU TỔNG HỢP: TRỌN BỘ BIỂU MẪU TỪ TRANG 8 ĐẾN TRANG 15 (8 TRANG CHUẨN)
# ==============================================================================
print("\n>>> 5. Đang tạo Trọn bộ Biểu mẫu (Trang 8-15): Tron_bo_Bieu_mau_Trang_8_den_15...")
doc_all = create_base_document(orientation="PORTRAIT", left_mm=30, right_mm=18, top_mm=20, bottom_mm=20)

# --- MẪU 1 (Trang 8-9) ---
add_p(doc_all, "PHỤ LỤC II. BIỂU MẪU", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1, line_spacing=1.15)
add_p(doc_all, "Mẫu I.01-ĐHNV", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=4, line_spacing=1.15)

top_t1 = doc_all.add_table(rows=1, cols=2)
top_t1.alignment = WD_TABLE_ALIGNMENT.CENTER
top_t1.autofit = False
cl1, cr1 = top_t1.cell(0, 0), top_t1.cell(0, 1)
cl1.width, cr1.width = Mm(52), Mm(110)

set_cell_margins(cl1, top=0, bottom=0, left=0, right=0)
p1 = cl1.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before, p1.paragraph_format.space_after, p1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
r = p1.add_run("TÊN CƠ QUAN CHỦ QUẢN")
r.font.size = Pt(12)

p2 = cl1.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before, p2.paragraph_format.space_after, p2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = p2.add_run("TÊN TỔ CHỨC")
r.font.size, r.font.bold = Pt(12), True

p3 = cl1.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before, p3.paragraph_format.space_after, p3.paragraph_format.line_spacing = Pt(0), Pt(0), 1.0
r = p3.add_run("______")
r.font.size, r.font.bold = Pt(11), True

set_cell_margins(cr1, top=0, bottom=0, left=0, right=0)
pr1 = cr1.paragraphs[0]
pr1.alignment = WD_ALIGN_PARAGRAPH.CENTER
pr1.paragraph_format.space_before, pr1.paragraph_format.space_after, pr1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
r = pr1.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
r.font.size, r.font.bold = Pt(12), True

pr2 = cr1.add_paragraph()
pr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pr2.paragraph_format.space_before, pr2.paragraph_format.space_after, pr2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = pr2.add_run("Độc lập - Tự do - Hạnh phúc")
r.font.size, r.font.bold = Pt(13), True

pr3 = cr1.add_paragraph()
pr3.alignment = WD_ALIGN_PARAGRAPH.CENTER
pr3.paragraph_format.space_before, pr3.paragraph_format.space_after, pr3.paragraph_format.line_spacing = Pt(0), Pt(3), 1.0
r = pr3.add_run("_________________")
r.font.size, r.font.bold = Pt(11), True

prd = cr1.add_paragraph()
prd.alignment = WD_ALIGN_PARAGRAPH.CENTER
prd.paragraph_format.space_before, prd.paragraph_format.space_after, prd.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r_d1 = prd.add_run("........")
r_d1.font.size = Pt(13)
r_d2 = prd.add_run(", ngày… tháng…  năm 20....")
r_d2.font.size, r_d2.font.italic = Pt(13), True

p_tit1 = add_p(doc_all, "", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=14, space_after=14, line_spacing=1.15)
r_t = p_tit1.add_run("PHIẾU ĐẶT HÀNG NHIỆM VỤ KHOA HỌC VÀ CÔNG NGHỆ")
r_t.font.size, r_t.font.bold = Pt(14), True
add_footnote_reference(p_tit1, footnote_id=1, font_size_half_pt=24)

# Page 8 items
add_p(doc_all, "1. Tên nhiệm vụ:\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "\t", space_after=4, has_tab_dots=True)
add_p(doc_all, "2. Thuộc lĩnh vực:", space_after=2)

tbl_f_all = doc_all.add_table(rows=4, cols=2)
tbl_f_all.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_f_all.autofit = False
for row_idx, (c1, c2) in enumerate(fields_data):
    cell1, cell2 = tbl_f_all.cell(row_idx, 0), tbl_f_all.cell(row_idx, 1)
    cell1.width, cell2.width = Mm(81), Mm(81)
    for cell, txt in [(cell1, c1), (cell2, c2)]:
        set_cell_margins(cell, top=0, bottom=0, left=50, right=50)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(2), 1.2
        p.paragraph_format.tab_stops.add_tab_stop(Mm(76), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        if txt:
            r1 = p.add_run(txt + "\t")
            r1.font.size = Pt(14)
            r2 = p.add_run("☐")
            r2.font.name, r2.font.size = 'Segoe UI Symbol', Pt(14)

add_p(doc_all, "3. Thuộc loại hình nhiệm vụ", space_before=4, space_after=2)
tbl_t_all = doc_all.add_table(rows=2, cols=2)
tbl_t_all.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_t_all.autofit = False
for row_idx, (c1, c2) in enumerate(types_data):
    cell1, cell2 = tbl_t_all.cell(row_idx, 0), tbl_t_all.cell(row_idx, 1)
    cell1.width, cell2.width = Mm(81), Mm(81)
    for cell, txt in [(cell1, c1), (cell2, c2)]:
        set_cell_margins(cell, top=0, bottom=0, left=50, right=50)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(2), 1.2
        p.paragraph_format.tab_stops.add_tab_stop(Mm(76), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        r1 = p.add_run(txt + "\t")
        r1.font.size = Pt(14)
        r2 = p.add_run("☐")
        r2.font.name, r2.font.size = 'Segoe UI Symbol', Pt(14)

add_p(doc_all, "4. Căn cứ đặt hàng nhiệm vụ:\t", space_before=4, space_after=2, has_tab_dots=True)
add_p(doc_all, "\t", space_after=4, has_tab_dots=True)
add_p(doc_all, "5. Tính cấp thiết của nhiệm vụ:\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "\t", space_after=4, has_tab_dots=True)
add_p(doc_all, "6. Mục tiêu:\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "\t", space_after=4, has_tab_dots=True)
add_p(doc_all, "7. Dự kiến các nội dung chính cần thực hiện:\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "\t", space_after=4, has_tab_dots=True)
add_p(doc_all, "8. Dự kiến các kết quả thực hiện nhiệm vụ và các chỉ tiêu cần đạt:\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "\t", space_after=4, has_tab_dots=True)
add_p(doc_all, "9. Hiệu quả và tác động của kết quả thực hiện nhiệm vụ:\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "\t", space_after=4, has_tab_dots=True)

# Page 9 items: ĐÚNG THỨ TỰ BẢN GỐC (10 -> 11 -> - Kinh phí NSNN -> - Kinh phí nguồn khác)
add_p(doc_all, "10.  Dự kiến kinh phí thực hiện:\t", space_after=2, keep_with_next=True, has_tab_dots=True)
add_p(doc_all, "11.  Dự kiến thời gian thực hiện:\t", space_after=2, keep_with_next=True, has_tab_dots=True)
add_p(doc_all, "- Kinh phí hỗ trợ từ NSNN:\t", space_after=2, keep_with_next=True, has_tab_dots=True)
add_p(doc_all, "- Kinh phí từ nguồn khác:\t", space_after=3, has_tab_dots=True)

add_p(doc_all, "12. Đề xuất tổ chức được xét giao trực tiếp (nếu có):\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "\t", space_after=4, has_tab_dots=True)

add_p(doc_all, "13. Tiếp nhận và phương án tổ chức quản lý, sử dụng kết quả của nhiệm vụ của", space_after=1)
add_p(doc_all, "cơ quan đặt hàng: ", space_after=2)
add_p(doc_all, "(Nêu rõ việc Nhà nước sẽ nắm giữ quyền quản lý, sử dụng, quyền sở hữu kết quả trong trường hợp nhà nước có yêu cầu tiếp nhận kết quả)", italic=True, space_after=2)
add_p(doc_all, "\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "\t", space_after=4, has_tab_dots=True)

add_p(doc_all, "14. Danh mục tài liệu tham khảo:", space_after=4)

add_p(doc_all, "15. Thông tin liên hệ:", space_after=2)
add_p(doc_all, "Tên tổ chức:\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "Đại diện tổ chức:\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "Mã định danh điện tử của tổ chức:\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "Điện thoại:\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "Email:\t", space_after=2, has_tab_dots=True)
add_p(doc_all, "Địa chỉ liên hệ:\t", space_after=14, has_tab_dots=True)

tbl_s_all1 = doc_all.add_table(rows=1, cols=2)
tbl_s_all1.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_s_all1.autofit = False
csl1, csr1 = tbl_s_all1.cell(0, 0), tbl_s_all1.cell(0, 1)
csl1.width, csr1.width = Mm(70), Mm(92)
set_cell_margins(csr1, top=0, bottom=0, left=0, right=0)
ps1 = csr1.paragraphs[0]
ps1.alignment = WD_ALIGN_PARAGRAPH.CENTER
ps1.paragraph_format.space_before, ps1.paragraph_format.space_after, ps1.paragraph_format.line_spacing = Pt(8), Pt(2), 1.15
r = ps1.add_run("ĐẠI DIỆN TỔ CHỨC ĐẶT HÀNG")
r.font.size, r.font.bold = Pt(14), True

ps2 = csr1.add_paragraph()
ps2.alignment = WD_ALIGN_PARAGRAPH.CENTER
ps2.paragraph_format.space_before, ps2.paragraph_format.space_after, ps2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = ps2.add_run("(Họ, tên và chữ ký - đóng dấu)")
r.font.size, r.font.italic = Pt(14), True


# --- MẪU 2 (Trang 10-11) ---
doc_all.add_page_break()
add_p(doc_all, "Mẫu 02 -TTMH", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=4, line_spacing=1.15)

top_t2 = doc_all.add_table(rows=1, cols=2)
top_t2.alignment = WD_TABLE_ALIGNMENT.CENTER
top_t2.autofit = False
cl2, cr2 = top_t2.cell(0, 0), top_t2.cell(0, 1)
cl2.width, cr2.width = Mm(65), Mm(97)

set_cell_margins(cl2, top=0, bottom=0, left=0, right=0)
ptl1 = cl2.paragraphs[0]
ptl1.alignment = WD_ALIGN_PARAGRAPH.CENTER
ptl1.paragraph_format.space_before, ptl1.paragraph_format.space_after, ptl1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
r = ptl1.add_run("CƠ QUAN TRÌNH")
r.font.size = Pt(12)

ptl2 = cl2.add_paragraph()
ptl2.alignment = WD_ALIGN_PARAGRAPH.CENTER
ptl2.paragraph_format.space_before, ptl2.paragraph_format.space_after, ptl2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = ptl2.add_run("PHÊ DUYỆT")
r.font.size, r.font.bold = Pt(12), True

ptl3 = cl2.add_paragraph()
ptl3.alignment = WD_ALIGN_PARAGRAPH.CENTER
ptl3.paragraph_format.space_before, ptl3.paragraph_format.space_after, ptl3.paragraph_format.line_spacing = Pt(0), Pt(2), 1.0
r = ptl3.add_run("-------")
r.font.size, r.font.bold = Pt(11), True

pnum2 = cl2.add_paragraph()
pnum2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pnum2.paragraph_format.space_before, pnum2.paragraph_format.space_after, pnum2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = pnum2.add_run("Số:……./TTr-…...")
r.font.size = Pt(13)

set_cell_margins(cr2, top=0, bottom=0, left=0, right=0)
ptr1 = cr2.paragraphs[0]
ptr1.alignment = WD_ALIGN_PARAGRAPH.CENTER
ptr1.paragraph_format.space_before, ptr1.paragraph_format.space_after, ptr1.paragraph_format.line_spacing = Pt(0), Pt(1), 1.15
r = ptr1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
r.font.size, r.font.bold = Pt(12), True

ptr2 = cr2.add_paragraph()
ptr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
ptr2.paragraph_format.space_before, ptr2.paragraph_format.space_after, ptr2.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
r = ptr2.add_run("Độc lập - Tự do - Hạnh phúc")
r.font.size, r.font.bold = Pt(13), True

pdate2 = cr2.add_paragraph()
pdate2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pdate2.paragraph_format.space_before, pdate2.paragraph_format.space_after, pdate2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = pdate2.add_run("……., ngày... tháng... năm…..")
r.font.size, r.font.italic = Pt(13), True

add_p(doc_all, "TỜ TRÌNH", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=2, line_spacing=1.15)
add_p(doc_all, "Về việc phê duyệt, hỗ trợ nhiệm vụ ứng dụng, chuyển giao tiến bộ\nkhoa học và công nghệ vào đời sống và sản xuất", 
      bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1, line_spacing=1.2)

p_subj_line_all2 = add_p(doc_all, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=10, line_spacing=1.0)
r_line_all2 = p_subj_line_all2.add_run("____________________")
r_line_all2.font.name = 'Times New Roman'
r_line_all2.font.size = Pt(11)
r_line_all2.font.bold = True

add_p(doc_all, "Kính gửi:   ", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=2, left_indent_mm=34)
add_p(doc_all, "- Sở Khoa học và Công nghệ;", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=1, left_indent_mm=56)
add_p(doc_all, "- Sở Tài chính.", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, left_indent_mm=56)

add_p(doc_all, "Căn cứ Luật Ngân sách nhà nước ngày 25 tháng 6 năm 2025;", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "Căn cứ Luật Quản lý, sử dụng tài sản công ngày 21 tháng 6 năm 2017;", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "Căn cứ Nghị định số 104/2026/NĐ-CP ngày 31 tháng 3 năm 2026 của Chính phủ quy định việc lập dự toán, quản lý, sử dụng và quyết toán chi thường xuyên để thực hiện các nhiệm vụ quy định tại Điều 40 Luật Ngân sách nhà nước;", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "Căn cứ các căn cứ pháp lý, văn bản liên quan khác (nếu có);", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=12.7)

p_intro_all2 = add_p(doc_all, "", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=12.7)
r1 = p_intro_all2.add_run("(Cơ quan trình phê duyệt)")
r1.font.name, r1.font.size, r1.font.italic = 'Times New Roman', Pt(14), True
r2 = p_intro_all2.add_run("…. Sở Khoa học và Công nghệ, Sở Tài chính xem xét thẩm định, phê duyệt hỗ trợ nhiệm vụ ứng dụng, chuyển giao tiến bộ khoa học và công nghệ vào đời sống và sản xuất năm 2027 cho ")
r2.font.name, r2.font.size = 'Times New Roman', Pt(14)
r3 = p_intro_all2.add_run("(Cơ quan trình phê duyệt)")
r3.font.name, r3.font.size, r3.font.italic = 'Times New Roman', Pt(14), True
r4 = p_intro_all2.add_run("….với các nội dung chủ yếu như sau:")
r4.font.name, r4.font.size = 'Times New Roman', Pt(14)

add_p(doc_all, "I. THÔNG TIN CHUNG", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=2)
add_p(doc_all, "1. Tên cơ quan, đơn vị đề xuất: ", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
add_p(doc_all, "2. Tên nhiệm vụ:", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
add_p(doc_all, "3. Dự kiến kinh phí:…. triệu đồng. Trong đó:", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
add_p(doc_all, "- Nguồn ngân sách nhà nước: .... triệu đồng", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "- Nguồn khác (nếu có): .... triệu đồng", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)

p_time_all = add_p(doc_all, "4. Thời gian thực hiện nhiệm vụ", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
add_footnote_reference(p_time_all, footnote_id=2, font_size_half_pt=22)
r_dot = p_time_all.add_run(".")
r_dot.font.size = Pt(14)

add_p(doc_all, "5. Các nội dung khác (nếu có).", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

add_p(doc_all, "II. NGUỒN KINH PHÍ THỰC HIỆN", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=2)
add_p(doc_all, "Đề nghị hỗ trợ từ nguồn ngân sách sự nghiệp KHCN, ĐMST&CĐS năm 2027 của tỉnh Hưng Yên", 
      font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=12.7)

add_p(doc_all, "III. HỒ SƠ TÀI LIỆU KÈM THEO", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=2)
add_p(doc_all, "- Thuyết minh chi tiết nhiệm vụ kèm theo dự toán;", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "- Báo giá và các tài liệu liên quan khác kèm theo.", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=12.7)

p_outro_all2 = add_p(doc_all, "", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=4, space_after=12, first_line_indent_mm=12.7)
r_o1 = p_outro_all2.add_run("(Cơ quan trình phê duyệt)")
r_o1.font.name, r_o1.font.size, r_o1.font.italic = 'Times New Roman', Pt(14), True
r_o2 = p_outro_all2.add_run("…. kính trình Sở Khoa học và Công nghệ, Sở Tài chính xem xét thẩm định, trình UBND tỉnh phê duyệt./.")
r_o2.font.name, r_o2.font.size = 'Times New Roman', Pt(14)

tbl_s_all2 = doc_all.add_table(rows=1, cols=2)
tbl_s_all2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_s_all2.autofit = False
cnn, ckt = tbl_s_all2.cell(0, 0), tbl_s_all2.cell(0, 1)
cnn.width, ckt.width = Mm(75), Mm(87)

set_cell_margins(cnn, top=0, bottom=0, left=0, right=0)
pnn = cnn.paragraphs[0]
pnn.alignment = WD_ALIGN_PARAGRAPH.LEFT
pnn.paragraph_format.space_before, pnn.paragraph_format.space_after, pnn.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
r = pnn.add_run("Nơi nhận:\n")
r.font.size, r.font.bold, r.font.italic = Pt(12), True, True
for line in ["- Như trên;", "- Các cơ quan có liên quan;", "- Lưu:..."]:
    r = pnn.add_run(line + "\n")
    r.font.size = Pt(11)

set_cell_margins(ckt, top=0, bottom=0, left=0, right=0)
pk1 = ckt.paragraphs[0]
pk1.alignment = WD_ALIGN_PARAGRAPH.CENTER
pk1.paragraph_format.space_before, pk1.paragraph_format.space_after, pk1.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
r = pk1.add_run("CƠ QUAN TRÌNH PHÊ DUYỆT")
r.font.size, r.font.bold = Pt(14), True

pk2 = ckt.add_paragraph()
pk2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pk2.paragraph_format.space_before, pk2.paragraph_format.space_after, pk2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = pk2.add_run("(Ký, ghi rõ họ tên, chức vụ và\nđóng dấu)")
r.font.size, r.font.italic = Pt(13), True


# --- MẪU 3 (Trang 12-14) ---
doc_all.add_page_break()
add_p(doc_all, "Mẫu 02 -TMMH", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=14, line_spacing=1.15)
add_p(doc_all, "………..(ĐƠN VỊ QUẢN LÝ)", bold=True, font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2, line_spacing=1.15)
add_p(doc_all, "………….. (ĐƠN VỊ THỰC HIỆN)", bold=True, font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=40, line_spacing=1.15)

add_p(doc_all, "THUYẾT MINH NHIỆM VỤ", bold=True, font_size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4, line_spacing=1.2)
add_p(doc_all, "ỨNG DỤNG, NHÂN RỘNG KẾT QUẢ NGHIÊN CỨU\nKHOA HỌC VÀ CÔNG NGHỆ", bold=True, font_size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=40, line_spacing=1.25)

add_p(doc_all, "Tên nhiệm vụ:……………………………………………………………….", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, line_spacing=1.2)
add_p(doc_all, "Ứng dụng, nhân rộng kết quả của nhiệm vụ KH&CN cấp tỉnh:", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=3, line_spacing=1.2)
add_p(doc_all, "“………..”", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6, line_spacing=1.2)
add_p(doc_all, "Đơn vị quản lý:…………", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, line_spacing=1.2)
add_p(doc_all, "Đơn vị thực hiện:………………", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, line_spacing=1.2)
add_p(doc_all, "Thời gian thực hiện: Năm 2027", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=50, line_spacing=1.2)
add_p(doc_all, "Hưng Yên, năm….", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0, line_spacing=1.2)

doc_all.add_page_break()
add_p(doc_all, "THUYẾT MINH NHIỆM VỤ", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2, line_spacing=1.15)
add_p(doc_all, "ỨNG DỤNG, NHÂN RỘNG KẾT QUẢ CỦA NHIỆM VỤ KH&CN", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=10, line_spacing=1.15)

add_p(doc_all, "I. THÔNG TIN CHUNG", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=3)
add_p(doc_all, "1. Tên nhiệm vụ:", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc_all, "- Ứng dụng, nhân rộng kết quả của nhiệm vụ KH&CN cấp tỉnh năm ………: “………………………………………………………………………..…..”", font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
add_p(doc_all, "Mã số (nếu có): ……………………………..", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

add_p(doc_all, "2. Tính cấp thiết cần triển khai thực hiện nhiệm vụ", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc_all, "Lý giải vì sao phải thực hiện nhiệm vụ này, sự cần thiết phải triển khai. Phân tích thực trạng sản xuất trong lĩnh vực đó và những hiệu quả trong thời gian vừa qua. Nhiệm vụ được triển khai sẽ góp phần nâng cao hiệu quả sản xuất, thay đổi tập quán, phương thức canh tác, chuẩn hóa các quy trình sản xuất….", 
      italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line_indent_mm=12.7)

add_p(doc_all, "3. Quy mô, địa điểm, thời gian thực hiện", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc_all, "- Quy mô: Diện tích, số lượng…", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "- Địa điểm:", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "- Thời gian: … tháng, từ tháng… đến tháng…năm …", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_line_indent_mm=12.7)

add_p(doc_all, "4. Kinh phí thực hiện", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc_all, "Tổng kinh phí:…. Triệu đồng", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "Trong đó:", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "- Ngân sách SNKH:…. Triệu đồng", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc_all, "- Nguồn khác: …. Triệu đồng", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_line_indent_mm=20)

add_p(doc_all, "5. Đơn vị phối hợp thực hiện, đơn vị hướng dẫn kỹ thuật", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
add_p(doc_all, "6. Đơn vị quản lý", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

add_p(doc_all, "II. MỤC TIÊU, NỘI DUNG, PHƯƠNG ÁN TỔ CHỨC VÀ TIẾN ĐỘ THỰC HIỆN", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=3)
add_p(doc_all, "1. Mục tiêu", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc_all, "1.1. Mục tiêu chung", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "1.2. Mục tiêu cụ thể", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "- Mục tiêu cụ thể 1", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc_all, "- Mục tiêu cụ thể 2", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc_all, "- Mục tiêu cụ thể 3", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_line_indent_mm=20)

add_p(doc_all, "2. Nội dung thực hiện (Luận giải những công việc cần thực hiện để hoàn thành các mục tiêu cụ thể, mục tiêu chung đã đề ra ở trên, đảm bảo tính thực tiễn, phù hợp với mục tiêu)", 
      italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3)
add_p(doc_all, "2.1. Nội dung 1", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "- Công việc 1", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc_all, "- Công việc 2", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc_all, "- Công việc 3", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3, first_line_indent_mm=20)

add_p(doc_all, "2.2. Nội dung 2", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "- Công việc 1", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc_all, "- Công việc 2", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc_all, "- Công việc 3", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3, first_line_indent_mm=20)

add_p(doc_all, "2.3. Nội dung 3", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "- Công việc 1", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc_all, "- Công việc 2", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc_all, "- Công việc 3", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=20)
add_p(doc_all, "……………", font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_line_indent_mm=12.7)

add_p(doc_all, "3. Phương án tổ chức và tiến độ thực hiện", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc_all, "3.1. Phương án tổ chức", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, first_line_indent_mm=12.7)
add_p(doc_all, "3.2. Tiến độ thực hiện", bold=True, italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_line_indent_mm=12.7)

tbl_s_sched = doc_all.add_table(rows=5, cols=6)
tbl_s_sched.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_s_sched.autofit = False
set_table_borders(tbl_s_sched)

for c_idx, (h_text, w) in enumerate(zip(headers_sched, col_widths_sched)):
    cell = tbl_s_sched.cell(0, c_idx)
    cell.width = w
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=80, bottom=80, left=60, right=60)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.1
    r = p.add_run(h_text)
    r.font.size, r.font.bold = Pt(13), True

for r_idx in range(1, 5):
    for c_idx, w in enumerate(col_widths_sched):
        cell = tbl_s_sched.cell(r_idx, c_idx)
        cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=100, bottom=100, left=60, right=60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 3, 5] else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), Pt(0), 1.1
        if c_idx == 0:
            r = p.add_run(str(r_idx))
            r.font.size = Pt(13)

add_p(doc_all, "III. DỰ KIẾN SẢN PHẨM, PHÂN TÍCH HIỆU QUẢ, KẾ HOẠCH NHÂN RỘNG", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=3)
add_p(doc_all, "1. Dự kiến sản phẩm của nhiệm vụ", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc_all, "2. Phân tích hiệu quả (hiệu quả kinh tế, xã hội, môi trường…)", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_p(doc_all, "3. Kế hoạch triển khai nhân rộng nhiệm vụ", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

add_p(doc_all, "Hưng Yên, ngày …. tháng…. năm …..", italic=True, font_size=14, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=4, space_after=8)

tbl_s_all3 = doc_all.add_table(rows=1, cols=2)
tbl_s_all3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_s_all3.autofit = False
cth, cql = tbl_s_all3.cell(0, 0), tbl_s_all3.cell(0, 1)
cth.width, cql.width = Mm(81), Mm(81)

set_cell_margins(cth, top=0, bottom=0, left=0, right=0)
pth1 = cth.paragraphs[0]
pth1.alignment = WD_ALIGN_PARAGRAPH.CENTER
pth1.paragraph_format.space_before, pth1.paragraph_format.space_after, pth1.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
r = pth1.add_run("ĐƠN VỊ THỰC HIỆN")
r.font.size, r.font.bold = Pt(14), True

pth2 = cth.add_paragraph()
pth2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pth2.paragraph_format.space_before, pth2.paragraph_format.space_after, pth2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = pth2.add_run("(Ký và đóng dấu)")
r.font.size, r.font.italic = Pt(14), True

set_cell_margins(cql, top=0, bottom=0, left=0, right=0)
pql1 = cql.paragraphs[0]
pql1.alignment = WD_ALIGN_PARAGRAPH.CENTER
pql1.paragraph_format.space_before, pql1.paragraph_format.space_after, pql1.paragraph_format.line_spacing = Pt(0), Pt(2), 1.15
r = pql1.add_run("ĐƠN VỊ QUẢN LÝ")
r.font.size, r.font.bold = Pt(14), True

pql2 = cql.add_paragraph()
pql2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pql2.paragraph_format.space_before, pql2.paragraph_format.space_after, pql2.paragraph_format.line_spacing = Pt(0), Pt(0), 1.15
r = pql2.add_run("(Ký và đóng dấu)")
r.font.size, r.font.italic = Pt(14), True


# --- MẪU 4 (Trang 15 - Landscape Section 1 Trang) ---
sec_land = doc_all.add_section(WD_SECTION_START.NEW_PAGE)
sec_land.page_width = Mm(297)
sec_land.page_height = Mm(210)
sec_land.orientation = WD_ORIENT.LANDSCAPE
sec_land.left_margin = Mm(18)
sec_land.right_margin = Mm(15)
sec_land.top_margin = Mm(15)
sec_land.bottom_margin = Mm(15)

build_landscape_phuluc(doc_all)


temp_docx_all = os.path.join(OUT_DIR, "temp_all.docx")
final_docx_all = os.path.join(OUT_DIR, "Tron_bo_Bieu_mau_Trang_8_den_15.docx")
final_pdf_all = os.path.join(OUT_DIR, "Tron_bo_Bieu_mau_Trang_8_den_15.pdf")

doc_all.save(temp_docx_all)
inject_footnotes_into_docx(temp_docx_all, final_docx_all, {
    1: "Phiếu đặt hàng được trình bày không quá 10 trang giấy khổ A4.",
    2: "Trong trường hợp nhiệm vụ có thời gian thực hiện trên 01 năm thì cần xác định dự toán kinh phí thực hiện trong từng năm."
})
os.remove(temp_docx_all)
print(f"  [DOCX Saved] {final_docx_all}")
export_docx_to_pdf(final_docx_all, final_pdf_all)

print("\n================================================================================")
print("HOÀN TẤT SINH TOÀN BỘ FILE DOCX VÀ PDF VÀO THƯ MỤC CHÍNH THỨC!")
print("================================================================================")
